"""
Streamlit dashboard for taxi chat information events.

Run locally:
    python -m streamlit run src/app.py -- --data-dir data/processed --db-path data/manual_actions.sqlite
"""

from __future__ import annotations

import argparse
import hashlib
import os
import re
from io import BytesIO
from collections import Counter
from datetime import date, datetime
from pathlib import Path

import numpy as np
import pandas as pd
import streamlit as st

from io_utils import read_table, read_source_csv
from manual_db import (
    connect,
    get_event_overrides,
    get_event_merges,
    get_message_overrides,
    get_message_exclusions,
    get_manual_events,
    get_message_topic_overrides,
    create_manual_event,
    save_event_override,
    merge_events,
    move_message,
    hide_message,
    mark_message_irrelevant,
    restore_message_relevance,
    get_key_message_pins,
    pin_key_message,
    unpin_key_message,
    save_message_topic_override,
    get_dashboard_summary,
    save_dashboard_summary,
)
from settings import STATUS_OPTIONS
from preprocess import run_preprocess, run_preprocess_from_dataframe
from persistent_store import (
    supabase_configured,
    list_periods,
    load_generated_tables_from_supabase,
    save_processed_tables_from_dir,
    make_period_id,
    save_uploaded_csv_to_storage,
    update_period_metadata,
    set_period_status,
    delete_period,
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(add_help=False)
    parser.add_argument("--data-dir", default=os.getenv("DASHBOARD_DATA_DIR", "data/processed"))
    parser.add_argument("--db-path", default=os.getenv("DASHBOARD_DB_PATH", "data/manual_actions.sqlite"))
    parser.add_argument("--upload-dir", default=os.getenv("DASHBOARD_UPLOAD_DIR", "data/uploads"))
    args, _ = parser.parse_known_args()
    return args


def get_secret_value(name: str, default: str = "") -> str:
    """Read Streamlit Secret first, then environment variable. Never expose values in UI."""
    try:
        value = st.secrets.get(name, default)
    except Exception:
        value = os.getenv(name, default)
    return str(value or "").strip()


def render_admin_mode() -> bool:
    """Return True when user is allowed to upload files and edit/moderate data."""
    admin_password = get_secret_value("ADMIN_PASSWORD", "")

    if "is_admin" not in st.session_state:
        st.session_state["is_admin"] = False

    if not admin_password:
        st.sidebar.warning(
            "ADMIN_PASSWORD не настроен. Режим загрузки и правок пока доступен всем пользователям со ссылкой."
        )
        return True

    if st.session_state.get("is_admin"):
        st.sidebar.success("Режим: администратор")
        if st.sidebar.button("Выйти из режима правок"):
            st.session_state["is_admin"] = False
            st.rerun()
        return True

    st.sidebar.info("Режим: просмотр")
    with st.sidebar.expander("Вход администратора", expanded=False):
        entered = st.text_input("Пароль администратора", type="password", key="admin_password_input")
        if st.button("Войти", key="admin_login_button"):
            if entered == admin_password:
                st.session_state["is_admin"] = True
                st.success("Режим администратора включен.")
                st.rerun()
            else:
                st.error("Неверный пароль.")

    return False


@st.cache_data(show_spinner=False)
def load_generated_tables(data_dir: str):
    events = read_table(data_dir, "events")
    discussions = read_table(data_dir, "discussions")
    messages = read_table(data_dir, "messages")
    discussion_messages = read_table(data_dir, "discussion_messages")
    event_discussions = read_table(data_dir, "event_discussions")

    for df, cols in [
        (events, ["start_date", "end_date"]),
        (discussions, ["start_date", "end_date"]),
        (messages, ["datetime"]),
    ]:
        for col in cols:
            if col in df.columns:
                df[col] = pd.to_datetime(df[col], errors="coerce")

    for col in ["is_hidden"]:
        if col in events.columns:
            events[col] = events[col].astype(str).str.lower().isin(["true", "1", "yes", "да"])

    for col in ["message_count", "chat_count", "author_count", "negative_count", "toxic_count", "discussion_count"]:
        if col in events.columns:
            events[col] = pd.to_numeric(events[col], errors="coerce").fillna(0).astype(int)

    for col in ["negative_share", "toxic_share", "importance_score"]:
        if col in events.columns:
            events[col] = pd.to_numeric(events[col], errors="coerce").fillna(0.0)

    return events, discussions, messages, discussion_messages, event_discussions


def resolve_merge_map(merges: pd.DataFrame) -> dict[str, str]:
    mapping = {}
    if merges is None or merges.empty:
        return mapping

    raw = dict(zip(merges["source_event_id"], merges["target_event_id"]))

    def resolve(x: str) -> str:
        seen = set()
        while x in raw and x not in seen:
            seen.add(x)
            x = raw[x]
        return x

    for source in raw:
        mapping[source] = resolve(source)
    return mapping


def normalize_title_for_auto_merge(title: str) -> str:
    """Normalize event title so visually identical topics are merged in the dashboard."""
    value = str(title or "").strip().lower().replace("ё", "е")
    value = re.sub(r"\s+", " ", value)
    value = re.sub(r"\s*[—–-]\s*волна\s*\d+\s*$", "", value)
    value = value.strip(" .,:;!?'\"«»()[]{}")
    return value

def normalize_manual_tags(tags: str) -> str:
    """Return tags in the same pipe-separated format as generated events."""
    raw = str(tags or "").replace(";", ",").replace("|", ",")
    items = []
    seen = set()
    for item in raw.split(","):
        item = re.sub(r"\s+", " ", item).strip()
        if not item:
            continue
        key = item.lower().replace("ё", "е")
        if key not in seen:
            seen.add(key)
            items.append(item)
    return "|".join(items)


def format_tags_for_display(tags: str) -> str:
    """Human-friendly tag rendering for message/event tables."""
    items = [
        re.sub(r"\s+", " ", t).strip()
        for t in re.split(r"[|;,]", str(tags or ""))
        if re.sub(r"\s+", " ", t).strip()
    ]
    seen = []
    for item in items:
        if item not in seen:
            seen.append(item)
    return "; ".join(seen)


def collect_available_source_topics(events: pd.DataFrame, messages: pd.DataFrame) -> list[str]:
    """Return source/main topics found in current data for manual assignment."""
    values: list[str] = []
    for df, cols in [
        (events, ["source_main_topics", "source_topics", "source_main_topic"]),
        (messages, ["source_main_topic", "source_topics"]),
    ]:
        if df is None or df.empty:
            continue
        for col in cols:
            if col not in df.columns:
                continue
            for raw in df[col].fillna("").astype(str):
                for item in re.split(r"[;|]", raw):
                    item = re.sub(r"\s+", " ", item).strip()
                    if item:
                        values.append(item)
    seen = {}
    for value in values:
        key = value.lower().replace("ё", "е")
        if key not in seen:
            seen[key] = value
    return sorted(seen.values(), key=lambda x: x.lower().replace("ё", "е"))


def message_topic_display(row: pd.Series, fallback_event_title: str = "") -> str:
    """Topic label shown in the full message feed."""
    for col in ["source_main_topic", "source_topics"]:
        value = str(row.get(col, "") or "").strip()
        if value and value.lower() not in {"nan", "none", "nat", "<na>"}:
            first = re.split(r"[;|]", value)[0].strip()
            if first:
                return first
    return str(fallback_event_title or "").strip()


def append_manual_events(events: pd.DataFrame, manual_events: pd.DataFrame) -> pd.DataFrame:
    """Append user-created information events to the generated event table."""
    if manual_events is None or manual_events.empty:
        return events

    events = events.copy()
    rows = []
    existing_ids = set(events.get("event_id", pd.Series(dtype=str)).astype(str))
    for _, row in manual_events.iterrows():
        event_id = str(row.get("event_id", "")).strip()
        if not event_id or event_id in existing_ids:
            continue
        rows.append({
            "event_id": event_id,
            "event_title": str(row.get("title", "")).strip() or "Новый инфоповод",
            "event_summary": str(row.get("summary", "")).strip(),
            "main_tag": "ручной",
            "main_tags": normalize_manual_tags(row.get("main_tags", "")),
            "keywords": "",
            "key_phrases": "",
            "start_date": pd.NaT,
            "end_date": pd.NaT,
            "discussion_count": 0,
            "message_count": 0,
            "chat_count": 0,
            "author_count": 0,
            "negative_count": 0,
            "toxic_count": 0,
            "negative_share": 0.0,
            "toxic_share": 0.0,
            "importance_score": 0.0,
            "status": str(row.get("status", "")).strip() or "новый",
            "is_hidden": str(row.get("hidden", "0")).lower() in ["1", "true", "yes", "да"],
            "is_manual": True,
        })

    if not rows:
        return events

    manual_df = pd.DataFrame(rows)
    for col in events.columns:
        if col not in manual_df.columns:
            manual_df[col] = np.nan
    for col in manual_df.columns:
        if col not in events.columns:
            events[col] = np.nan

    return pd.concat([events, manual_df[events.columns]], ignore_index=True)





SUMMARY_STOPWORDS = {
    "это", "как", "что", "или", "для", "при", "про", "без", "есть", "нет", "все", "уже", "еще", "ещё",
    "там", "тут", "они", "она", "его", "мне", "нам", "вам", "тебя", "себя", "так", "такой", "такая",
    "если", "когда", "куда", "где", "кто", "чем", "надо", "нужно", "можно", "нельзя", "будет", "были",
    "был", "была", "будут", "этот", "эта", "эти", "эту", "того", "тоже", "очень", "просто",
    "сейчас", "сегодня", "вчера", "завтра", "потом", "теперь", "больше", "меньше", "через", "после",
    "водитель", "водители", "такси", "чат", "чате", "сообщение", "сказал", "говорят", "пишут",
    "яндекс", "yandex", "про", "таксометр", "машина", "заказ", "заказы",
}

SUMMARY_RULES = [
    ("законопроект и новые требования к работе такси", [r"законопроект\w*", r"\bзакон\w*", r"регулирован\w*", r"требован\w*", r"минтранс", r"госдум", r"реестр", r"разрешени\w*"]),
    ("налоги, патенты и самозанятость водителей", [r"налог\w*", r"патент\w*", r"самозанят\w*", r"нпд", r"деклараци\w*", r"фнс", r"штраф\w*"]),
    ("повышение или снижение коэффициентов", [r"коэфф\w*", r"кэф\w*", r"повыш\w*.{0,40}коэфф\w*", r"сниз\w*.{0,40}коэфф\w*", r"вырос\w*.{0,40}коэфф\w*"]),
    ("приоритет, тарифы и распределение заказов", [r"приоритет\w*", r"тариф\w*", r"эконом", r"комфорт", r"комисси\w*", r"раздач\w* заказ\w*", r"распределен\w* заказ\w*"]),
    ("невозможность загрузить или установить обновление", [r"обновлен\w*", r"обновить", r"скачать", r"загруз\w*", r"установ\w*", r"верси\w*", r"апдейт"]),
    ("сбои, ошибки и зависания приложения", [r"сбой\w*", r"ошибк\w*", r"не работает", r"завис\w*", r"вылета\w*", r"лага\w*", r"глюч\w*", r"баг\w*"]),
    ("проблемы с заказами, отменами и назначением поездок", [r"отмен\w*", r"назнач\w*", r"принять заказ", r"заказ\w*.{0,40}не приход", r"не дает заказ", r"цепочк\w*", r"подач\w*"]),
    ("блокировки, доступ к аккаунту и проверки", [r"блокиров\w*", r"заблок\w*", r"аккаунт\w*", r"доступ\w*", r"провер\w*", r"верификац\w*", r"фотоконтроль"]),
    ("оплата, выплаты и удержания", [r"оплат\w*", r"выплат\w*", r"деньг\w*", r"баланс\w*", r"удерж\w*", r"перевод\w*", r"компенсац\w*"]),
    ("забастовка, бойкот и коллективные действия", [r"забастов\w*", r"бойкот\w*", r"стачк\w*", r"не выход\w* на лини\w*", r"акци\w* протест\w*"]),
    ("WB Такси, условия запуска и сравнение с агрегаторами", [r"wb\s*такси", r"wildberries", r"вайлдбер\w*", r"вб\s*такси", r"wb"]),
    ("Фастен и альтернативные сервисы для водителей", [r"фастен", r"fasten", r"альтернатив\w* агрегатор\w*", r"нов\w* сервис\w*"]),
    ("карты, адреса, геолокация и навигация", [r"карт\w*", r"адрес\w*", r"геолокац\w*", r"gps", r"навигатор", r"точк\w* подач\w*", r"маршрут\w*"]),
    ("аэропорты, очереди и заказы из аэропорта", [r"аэропорт\w*", r"шереметьево", r"домодедово", r"внуково", r"пулково", r"очеред\w*"]),
    ("детские кресла и требования к заказам с детьми", [r"детск\w* кресл\w*", r"кресл\w*", r"ребен\w*", r"ребён\w*", r"дет\w* тариф\w*"]),
]


def normalize_text_for_summary(text: str) -> str:
    text = str(text or "").lower().replace("ё", "е")
    text = re.sub(r"https?://\S+|t\.me/\S+", " ", text)
    text = re.sub(r"[^а-яa-z0-9\s-]", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def summary_tokens(text: str) -> list[str]:
    text = normalize_text_for_summary(text)
    tokens = re.findall(r"[а-яa-z0-9]{3,}", text)
    return [t for t in tokens if t not in SUMMARY_STOPWORDS and not t.isdigit()]


def top_readable_phrases(texts, top_n: int = 5) -> list[str]:
    """Return short readable frequent phrases as fallback details for summaries."""
    counter: dict[str, int] = {}
    for text in texts:
        tokens = summary_tokens(str(text)[:2500])
        for a, b in zip(tokens, tokens[1:]):
            if a == b:
                continue
            phrase = f"{a} {b}"
            counter[phrase] = counter.get(phrase, 0) + 1
    ranked = sorted(counter.items(), key=lambda x: x[1], reverse=True)
    return [p for p, c in ranked[:top_n] if c >= 2]


def extract_summary_mentions(event_title: str, main_tags: str, event_messages: pd.DataFrame, keywords: str = "", phrases: str = "") -> list[str]:
    """Extract user-facing thesis bullets from the most frequent signals in event messages."""
    text_series = event_messages.get("text_clean", pd.Series(dtype=str)).dropna().astype(str)
    if text_series.empty:
        base = [p.strip() for p in str(phrases or "").split("|") if p.strip()]
        base += [k.strip() for k in str(keywords or "").split("|") if k.strip()]
        return base[:5]

    full_text = "\n".join(text_series.head(350).tolist())
    normalized = normalize_text_for_summary(full_text)

    scored: list[tuple[int, str]] = []
    for label, patterns in SUMMARY_RULES:
        count = 0
        for pattern in patterns:
            count += len(re.findall(pattern, normalized, flags=re.IGNORECASE))
        if count > 0:
            scored.append((count, label))

    title_tags = normalize_text_for_summary(f"{event_title} {main_tags}")
    boosted = []
    for count, label in scored:
        boost = 2 if any(token in title_tags for token in summary_tokens(label)[:2]) else 0
        boosted.append((count + boost, label))

    mentions = [label for _, label in sorted(boosted, key=lambda x: x[0], reverse=True)]

    for phrase in top_readable_phrases(text_series, top_n=5):
        phrase = phrase.strip()
        if phrase and all(phrase not in m for m in mentions):
            mentions.append(phrase)
        if len(mentions) >= 6:
            break

    clean: list[str] = []
    seen: set[str] = set()
    for item in mentions:
        key = normalize_text_for_summary(item)
        if not key or key in seen:
            continue
        seen.add(key)
        clean.append(item)
        if len(clean) >= 6:
            break

    return clean


def build_event_description(event_title: str, main_tags: str, event_messages: pd.DataFrame, keywords: str = "", phrases: str = "") -> str:
    """Build concise thesis-style summary for table/card display."""
    mentions = extract_summary_mentions(event_title, main_tags, event_messages, keywords=keywords, phrases=phrases)
    if not mentions:
        return "В теме обсуждались связанные сообщения по выбранному инфоповоду."
    return "В теме обсуждались: " + "; ".join(mentions[:6]) + "."

def build_auto_title_merge_map(events: pd.DataFrame) -> dict[str, str]:
    """Map duplicate event titles to one representative event id.

    This is intentionally done at the dashboard layer: source data stays detailed,
    while the user sees one consolidated topic for identical event titles.
    """
    if events.empty or "final_event_id" not in events.columns or "event_title" not in events.columns:
        return {}

    work_events = events.copy()
    if "is_manual" in work_events.columns:
        work_events = work_events[~work_events["is_manual"].astype(str).str.lower().isin(["true", "1", "yes", "да"])]

    representatives = (
        work_events.groupby("final_event_id", as_index=False)
        .agg(
            event_title=("event_title", "first"),
            message_count=("message_count", "sum"),
            importance_score=("importance_score", "max"),
        )
    )
    representatives["title_key"] = representatives["event_title"].apply(normalize_title_for_auto_merge)
    representatives = representatives[representatives["title_key"].astype(str).str.len() > 0]

    title_merge_map: dict[str, str] = {}
    for _, group in representatives.groupby("title_key", sort=False):
        if len(group) <= 1:
            continue
        ordered = group.sort_values(["message_count", "importance_score"], ascending=False)
        target_id = str(ordered.iloc[0]["final_event_id"])
        for source_id in ordered["final_event_id"].astype(str):
            title_merge_map[source_id] = target_id

    return title_merge_map


def apply_manual_edits(
    events: pd.DataFrame,
    messages: pd.DataFrame,
    discussion_messages: pd.DataFrame,
    event_discussions: pd.DataFrame,
    overrides: pd.DataFrame,
    merges: pd.DataFrame,
    message_overrides: pd.DataFrame,
    message_exclusions: pd.DataFrame | None = None,
    manual_events: pd.DataFrame | None = None,
    message_topic_overrides: pd.DataFrame | None = None,
):
    events = append_manual_events(events.copy(), manual_events)
    links = event_discussions.copy()
    msg_links = discussion_messages.merge(links, on="discussion_id", how="left")

    merge_map = resolve_merge_map(merges)

    def apply_manual_event_merge(event_id):
        return merge_map.get(event_id, event_id)

    links["event_id"] = links["event_id"].apply(apply_manual_event_merge)
    msg_links["event_id"] = msg_links["event_id"].apply(apply_manual_event_merge)
    events["final_event_id"] = events["event_id"].apply(apply_manual_event_merge)

    auto_title_merge_map = build_auto_title_merge_map(events)

    def canonical_event_id(event_id):
        event_id = merge_map.get(event_id, event_id)
        return auto_title_merge_map.get(str(event_id), event_id)

    if auto_title_merge_map:
        links["event_id"] = links["event_id"].apply(canonical_event_id)
        msg_links["event_id"] = msg_links["event_id"].apply(canonical_event_id)
        events["final_event_id"] = events["final_event_id"].apply(canonical_event_id)

    if message_overrides is not None and not message_overrides.empty:
        move_map = {
            r["message_id"]: canonical_event_id(r["target_event_id"])
            for _, r in message_overrides.iterrows()
            if str(r.get("target_event_id", "")).strip()
        }
        if move_map:
            msg_links["event_id"] = msg_links.apply(
                lambda r: move_map.get(r["message_id"], r["event_id"]),
                axis=1,
            )

    if message_exclusions is not None and not message_exclusions.empty and len(msg_links):
        exclusions = message_exclusions.copy()
        exclusions["event_id"] = exclusions["event_id"].apply(canonical_event_id)
        excluded_pairs = set(
            zip(
                exclusions["message_id"].astype(str),
                exclusions["event_id"].astype(str),
            )
        )
        msg_links = msg_links[
            ~msg_links.apply(
                lambda r: (str(r.get("message_id", "")), str(r.get("event_id", ""))) in excluded_pairs,
                axis=1,
            )
        ]

    msg_event = (
        msg_links[["message_id", "event_id"]]
        .dropna()
        .drop_duplicates()
        .rename(columns={"event_id": "final_event_id"})
    )

    enriched_messages = messages.merge(msg_event, on="message_id", how="left")

    if message_topic_overrides is not None and not message_topic_overrides.empty and "message_id" in enriched_messages.columns:
        topic_overrides = message_topic_overrides.copy()
        topic_overrides["message_id"] = topic_overrides["message_id"].astype(str)
        for _, topic_row in topic_overrides.iterrows():
            msg_id = str(topic_row.get("message_id", "")).strip()
            if not msg_id:
                continue
            mask = enriched_messages["message_id"].astype(str) == msg_id
            if not mask.any():
                continue
            main_topic = str(topic_row.get("source_main_topic", "") or "").strip()
            source_topics = str(topic_row.get("source_topics", "") or "").strip() or main_topic
            if main_topic:
                if "source_main_topic" not in enriched_messages.columns:
                    enriched_messages["source_main_topic"] = ""
                enriched_messages.loc[mask, "source_main_topic"] = main_topic
            if source_topics:
                if "source_topics" not in enriched_messages.columns:
                    enriched_messages["source_topics"] = ""
                enriched_messages.loc[mask, "source_topics"] = source_topics

    if message_overrides is not None and not message_overrides.empty:
        hidden_msg = set(
            message_overrides.loc[
                message_overrides["hidden"].astype(str).isin(["1", "true", "True"]),
                "message_id",
            ]
        )
        enriched_messages["message_hidden"] = enriched_messages["message_id"].isin(hidden_msg)
    else:
        enriched_messages["message_hidden"] = False

    # If an uploaded file contains a human/external relevance column, exclude
    # non-relevant messages from generated info events and summaries without
    # deleting them from the stored message base.
    if "source_relevant" in enriched_messages.columns:
        relevant_mask = enriched_messages["source_relevant"].astype(str).str.lower().isin(["true", "1", "yes", "да"])
        # Empty values mean that the file did not provide relevance markup.
        empty_rel = enriched_messages["source_relevant"].astype(str).str.strip().eq("")
        enriched_messages["message_hidden"] = enriched_messages["message_hidden"].astype(bool) | (~relevant_mask & ~empty_rel)

    rows = []
    for final_id, group in events.groupby("final_event_id", sort=False):
        target = events[events["event_id"] == final_id]
        base = target.iloc[0] if len(target) else group.iloc[0]

        group_messages = enriched_messages[
            (enriched_messages["final_event_id"] == final_id) & (~enriched_messages["message_hidden"].astype(bool))
        ]

        all_tags = sorted(set(t for tags in group.get("main_tags", pd.Series(dtype=str)).fillna("") for t in str(tags).split("|") if t.strip()))
        if len(group_messages) and "tags" in group_messages.columns:
            message_tags = sorted(set(t for tags in group_messages["tags"].fillna("") for t in str(tags).split("|") if t.strip()))
            all_tags = sorted(set(all_tags) | set(message_tags))

        all_source_main_topics = sorted(set(
            t.strip()
            for topics in group.get("source_main_topic", pd.Series(dtype=str)).fillna("")
            for t in str(topics).split(";")
            if t.strip()
        ))
        if len(group_messages) and "source_main_topic" in group_messages.columns:
            msg_source_topics = sorted(set(
                t.strip()
                for topics in group_messages["source_main_topic"].fillna("")
                for t in str(topics).split(";")
                if t.strip()
            ))
            all_source_main_topics = sorted(set(all_source_main_topics) | set(msg_source_topics))

        all_source_topics = sorted(set(
            t.strip()
            for topics in group.get("source_topics", pd.Series(dtype=str)).fillna("")
            for t in re.split(r"[;|]", str(topics))
            if t.strip()
        ))
        if len(group_messages) and "source_topics" in group_messages.columns:
            msg_source_topics_all = sorted(set(
                t.strip()
                for topics in group_messages["source_topics"].fillna("")
                for t in re.split(r"[;|]", str(topics))
                if t.strip()
            ))
            all_source_topics = sorted(set(all_source_topics) | set(msg_source_topics_all))

        keywords = sorted(set(t for tags in group.get("keywords", pd.Series(dtype=str)).fillna("") for t in str(tags).split("|") if t.strip()))
        phrases = sorted(set(t for tags in group.get("key_phrases", pd.Series(dtype=str)).fillna("") for t in str(tags).split("|") if t.strip()))

        start_date = group_messages["datetime"].min() if "datetime" in group_messages and len(group_messages) else group["start_date"].min()
        end_date = group_messages["datetime"].max() if "datetime" in group_messages and len(group_messages) else group["end_date"].max()

        msg_count = int(group_messages["message_id"].nunique()) if len(group_messages) else int(group["message_count"].sum())
        chat_count = int(group_messages["chat_id"].nunique()) if "chat_id" in group_messages and len(group_messages) else int(group["chat_count"].sum())
        author_count = int(group_messages["author_id"].nunique()) if "author_id" in group_messages and len(group_messages) else int(group["author_count"].sum())
        negative_count = int(group_messages["is_negative"].astype(str).str.lower().isin(["true", "1"]).sum()) if "is_negative" in group_messages and len(group_messages) else int(group["negative_count"].sum())
        toxic_count = int(group_messages["is_toxic"].astype(str).str.lower().isin(["true", "1"]).sum()) if "is_toxic" in group_messages and len(group_messages) else int(group["toxic_count"].sum())

        event_title = str(base.get("event_title", ""))
        manual_summary = str(base.get("event_summary", "") or "").strip() if str(base.get("is_manual", "")).lower() in ["true", "1", "yes", "да"] else ""
        event_summary = manual_summary or build_event_description(
            event_title,
            "|".join(all_tags),
            group_messages,
            keywords="|".join(keywords),
            phrases="|".join(phrases),
        )

        source_event_ids = sorted({str(x) for x in group.get("event_id", pd.Series(dtype=str)).dropna().astype(str).tolist()})

        rows.append({
            "event_id": final_id,
            "source_event_ids": "|".join(source_event_ids),
            "source_event_count": len(source_event_ids),
            "event_title": event_title,
            "event_summary": event_summary,
            "main_tag": base.get("main_tag", ""),
            "main_tags": "|".join(all_tags),
            "source_main_topics": "; ".join(all_source_main_topics),
            "source_topics": "; ".join(all_source_topics),
            "keywords": "|".join(keywords),
            "key_phrases": "|".join(phrases),
            "start_date": start_date,
            "end_date": end_date,
            "discussion_count": int(links[links["event_id"] == final_id]["discussion_id"].nunique()) if len(links) else 0,
            "message_count": msg_count,
            "chat_count": chat_count,
            "author_count": author_count,
            "negative_count": negative_count,
            "toxic_count": toxic_count,
            "importance_score": float(group["importance_score"].max()),
            "status": base.get("status", "новый"),
            "is_hidden": bool(base.get("is_hidden", False)) if str(base.get("is_manual", "")).lower() in ["true", "1", "yes", "да"] else False,
            "is_manual": str(base.get("is_manual", "")).lower() in ["true", "1", "yes", "да"],
        })

    visible_events = pd.DataFrame(rows)

    if overrides is not None and not overrides.empty and len(visible_events):
        ov = overrides.set_index("event_id")
        for idx, row in visible_events.iterrows():
            event_id = row["event_id"]
            if event_id not in ov.index:
                continue
            o = ov.loc[event_id]
            if str(o.get("title", "")).strip():
                visible_events.at[idx, "event_title"] = o["title"]
            if str(o.get("summary", "")).strip():
                visible_events.at[idx, "event_summary"] = o["summary"]
            if str(o.get("status", "")).strip():
                visible_events.at[idx, "status"] = o["status"]
            if str(o.get("priority", "")).strip():
                visible_events.at[idx, "priority"] = o["priority"]
            hidden_val = str(o.get("hidden", "0")).lower()
            visible_events.at[idx, "is_hidden"] = hidden_val in ["1", "true", "yes", "да"]

    if len(visible_events):
        visible_events["negative_share"] = np.where(
            visible_events["message_count"] > 0,
            visible_events["negative_count"] / visible_events["message_count"],
            0,
        )
        visible_events["toxic_share"] = np.where(
            visible_events["message_count"] > 0,
            visible_events["toxic_count"] / visible_events["message_count"],
            0,
        )
        visible_events = visible_events.sort_values(["importance_score", "message_count"], ascending=False)

    return visible_events, enriched_messages


def format_pct(value) -> str:
    try:
        return f"{float(value) * 100:.0f}%"
    except Exception:
        return "0%"


def parse_date_value(value):
    """Parse a date value safely and prefer Russian DD.MM.YYYY for user input."""
    if value is None:
        return None
    try:
        if pd.isna(value):
            return None
    except (TypeError, ValueError):
        pass
    if isinstance(value, pd.Timestamp):
        return value.date()
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value

    text = str(value or "").strip()
    if not text or text.lower() in {"nat", "nan", "none", "null", "<na>"}:
        return None

    iso_match = re.match(r"^(\d{4})-(\d{1,2})-(\d{1,2})", text)
    if iso_match:
        try:
            return date(int(iso_match.group(1)), int(iso_match.group(2)), int(iso_match.group(3)))
        except ValueError:
            return None

    ru_match = re.match(r"^(\d{1,2})[.](\d{1,2})[.](\d{2,4})$", text)
    if ru_match:
        day, month, year = [int(x) for x in ru_match.groups()]
        if year < 100:
            year += 2000
        try:
            return date(year, month, day)
        except ValueError:
            return None

    dt = pd.to_datetime(text, errors="coerce", dayfirst=True)
    if pd.isna(dt):
        return None
    return dt.date()

def date_to_iso(value) -> str | None:
    parsed = parse_date_value(value)
    return parsed.isoformat() if parsed else None


def format_date(value) -> str:
    """Return a user-facing date without time: DD.MM.YYYY."""
    parsed = parse_date_value(value)
    if parsed is None:
        return ""
    try:
        if pd.isna(parsed):
            return ""
    except (TypeError, ValueError):
        pass
    try:
        return parsed.strftime("%d.%m.%Y")
    except Exception:
        return ""

def format_date_series(series: pd.Series) -> pd.Series:
    """Format pandas datetime/string series as DD.MM.YYYY strings for display tables."""
    return series.apply(format_date).fillna("")


def period_note_from_manifest(value) -> str:
    if isinstance(value, dict):
        return str(value.get("period_note", "") or "")
    return ""


def parse_user_date(value: str):
    return date_to_iso(value)


def format_period(row: pd.Series) -> str:
    start_s = format_date(row.get("start_date"))
    end_s = format_date(row.get("end_date"))

    if not start_s:
        return end_s
    if not end_s or start_s == end_s:
        return start_s
    return f"{start_s} — {end_s}"

def get_selected_rows(event) -> list[int]:
    try:
        return list(event.selection.rows)
    except Exception:
        try:
            return list(event.get("selection", {}).get("rows", []))
        except Exception:
            return []


def normalize_search_query(value: str) -> str:
    """Normalize user-entered text search query for Russian text."""
    return re.sub(r"\s+", " ", str(value or "").lower().replace("ё", "е")).strip()


def filter_messages_by_word(messages: pd.DataFrame, query: str) -> pd.DataFrame:
    """Return visible messages whose text contains the entered word or phrase."""
    q = normalize_search_query(query)
    if not q or "text_clean" not in messages.columns:
        return messages.iloc[0:0].copy()

    work = messages.copy()
    if "message_hidden" in work.columns:
        work = work[~work["message_hidden"].astype(bool)]

    text = (
        work["text_clean"]
        .fillna("")
        .astype(str)
        .str.lower()
        .str.replace("ё", "е", regex=False)
    )
    return work[text.str.contains(q, regex=False, na=False)].copy()


def apply_filters(events: pd.DataFrame, messages: pd.DataFrame) -> tuple[pd.DataFrame, str, pd.DataFrame]:
    st.sidebar.header("Фильтры")

    filtered = events.copy()
    empty_messages = messages.iloc[0:0].copy() if isinstance(messages, pd.DataFrame) else pd.DataFrame()
    if filtered.empty:
        return filtered, "", empty_messages

    word_query = st.sidebar.text_input(
        "Слово в тексте сообщений",
        placeholder="например: налог, коэффициент, обновление",
        help="Фильтр ищет введенное слово или фразу именно в текстах сообщений, а не в названии темы.",
    )
    word_matches = filter_messages_by_word(messages, word_query) if word_query else empty_messages

    all_tags = sorted(set(t for tags in filtered.get("main_tags", pd.Series(dtype=str)).fillna("") for t in str(tags).split("|") if t.strip()))
    selected_tags = st.sidebar.multiselect("Теги", all_tags)

    all_source_topics = sorted(set(
        t.strip()
        for topics in filtered.get("source_main_topics", pd.Series(dtype=str)).fillna("")
        for t in str(topics).split(";")
        if t.strip()
    ))
    selected_source_topics = st.sidebar.multiselect(
        "Тема из файла",
        all_source_topics,
        help="Появляется, если в загруженной выгрузке есть колонка «Основная тема». Используется как верхний тематический слой.",
    ) if all_source_topics else []

    statuses = [s for s in STATUS_OPTIONS if s in set(filtered.get("status", pd.Series(dtype=str)).fillna("").astype(str))]
    selected_statuses = st.sidebar.multiselect("Статус", statuses)

    only_attention = st.sidebar.checkbox("Только требующие внимания", value=False)

    with st.sidebar.expander("Дополнительно", expanded=False):
        show_hidden = st.checkbox("Показывать скрытые", value=False)
        negative_only = st.checkbox("Только с негативом", value=False)
        min_importance = st.slider(
            "Минимальная важность",
            0.0,
            float(max(filtered["importance_score"].max(), 1.0)),
            0.0,
        )

    if not show_hidden and "is_hidden" in filtered.columns:
        filtered = filtered[~filtered["is_hidden"].astype(bool)]

    if word_query:
        matched_event_ids = set(word_matches.get("final_event_id", pd.Series(dtype=str)).dropna().astype(str))
        filtered = filtered[filtered["event_id"].astype(str).isin(matched_event_ids)]

    if selected_tags:
        filtered = filtered[
            filtered["main_tags"].fillna("").apply(
                lambda x: bool(set(selected_tags) & {t.strip() for t in str(x).split("|") if t.strip()})
            )
        ]

    if selected_source_topics and "source_main_topics" in filtered.columns:
        selected_source_set = set(selected_source_topics)
        filtered = filtered[
            filtered["source_main_topics"].fillna("").apply(
                lambda x: bool(selected_source_set & {t.strip() for t in str(x).split(";") if t.strip()})
            )
        ]

    if selected_statuses:
        filtered = filtered[filtered["status"].isin(selected_statuses)]

    filtered = filtered[filtered["importance_score"] >= min_importance]

    if negative_only:
        filtered = filtered[filtered["negative_count"] > 0]

    if only_attention:
        filtered = filtered[
            (filtered["importance_score"] >= filtered["importance_score"].quantile(0.70))
            | (filtered["negative_share"] >= 0.25)
            | (filtered["toxic_share"] >= 0.10)
            | (filtered["main_tags"].fillna("").str.contains("Забастовка|Законы", regex=True))
        ]

    return filtered.sort_values(["importance_score", "message_count"], ascending=False), word_query, word_matches

def create_manual_event_form(conn, key_prefix: str = "manual_event", *, compact: bool = False) -> str | None:
    """Render a form for creating a user-defined information event."""
    with st.form(f"{key_prefix}_form"):
        title = st.text_input("Название инфоповода", placeholder="Например: Проблемы с детскими креслами")
        summary = st.text_area(
            "Описание",
            placeholder="Например: В теме обсуждались: требования к детским креслам; отказы от заказов с детьми; штрафы.",
            height=90 if compact else 130,
        )
        tags = st.text_input("Теги", placeholder="через запятую: Яндекс, детские кресла, тарифы")
        status = st.selectbox(
            "Статус",
            STATUS_OPTIONS,
            index=STATUS_OPTIONS.index("новый") if "новый" in STATUS_OPTIONS else 0,
            key=f"{key_prefix}_status",
        )
        note = st.text_input("Комментарий модератора", value="", key=f"{key_prefix}_note")
        submitted = st.form_submit_button("Создать инфоповод")
        if submitted:
            try:
                new_id = create_manual_event(
                    conn,
                    title=title,
                    summary=summary,
                    status=status,
                    main_tags=normalize_manual_tags(tags),
                    note=note,
                )
                st.success("Инфоповод создан. Теперь в него можно переносить сообщения.")
                st.cache_data.clear()
                return new_id
            except Exception as e:
                st.error(str(e))
    return None


def show_manual_event_creator(conn):
    with st.sidebar.expander("Создать инфоповод", expanded=False):
        st.caption("Используйте, если нужной темы нет в списке. Новый инфоповод появится в таблице и в списке для переноса сообщений.")
        new_id = create_manual_event_form(conn, key_prefix="sidebar_create_event", compact=True)
        if new_id:
            st.rerun()


def messages_for_events(messages: pd.DataFrame, events: pd.DataFrame) -> pd.DataFrame:
    """Return visible messages that belong to the currently displayed events."""
    if not isinstance(messages, pd.DataFrame) or messages.empty or not isinstance(events, pd.DataFrame) or events.empty:
        return pd.DataFrame()
    if "event_id" not in events.columns or "final_event_id" not in messages.columns:
        return pd.DataFrame()
    event_ids = set(events["event_id"].dropna().astype(str))
    result = messages[messages["final_event_id"].astype(str).isin(event_ids)].copy()
    if "message_hidden" in result.columns:
        result = result[~result["message_hidden"].astype(bool)]
    return result


def show_kpis(events: pd.DataFrame, messages: pd.DataFrame | None = None):
    visible_messages = messages.copy() if isinstance(messages, pd.DataFrame) else pd.DataFrame()
    if "message_hidden" in visible_messages.columns:
        visible_messages = visible_messages[~visible_messages["message_hidden"].astype(bool)]

    if not visible_messages.empty:
        message_count = int(visible_messages["message_id"].nunique()) if "message_id" in visible_messages.columns else int(len(visible_messages))
        chat_col = "chat_id" if "chat_id" in visible_messages.columns else "chat_title" if "chat_title" in visible_messages.columns else None
        chat_count = int(visible_messages[chat_col].nunique()) if chat_col else 0
        if "is_negative" in visible_messages.columns:
            negative_count = int(visible_messages["is_negative"].astype(str).str.lower().isin(["true", "1", "yes", "да"]).sum())
        elif "sentiment" in visible_messages.columns:
            negative_count = int(visible_messages["sentiment"].astype(str).str.lower().str.contains("нег|neg", regex=True, na=False).sum())
        else:
            negative_count = 0
    else:
        message_count = int(events["message_count"].sum()) if len(events) and "message_count" in events.columns else 0
        # Do not use max(chat_count): it undercounts multi-period selections.
        # Sum is only a fallback when message-level data is unavailable.
        chat_count = int(events["chat_count"].sum()) if len(events) and "chat_count" in events.columns else 0
        negative_count = int(events["negative_count"].sum()) if len(events) and "negative_count" in events.columns else 0

    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Инфоповодов", f"{len(events):,}".replace(",", " "))
    c2.metric("Сообщений", f"{message_count:,}".replace(",", " "))
    c3.metric("Чатов", f"{chat_count:,}".replace(",", " "))
    c4.metric("Негатив", format_pct(negative_count / message_count) if message_count else "0%")
    c5.metric("Высокая важность", int((events["importance_score"] >= events["importance_score"].quantile(0.75)).sum()) if len(events) and "importance_score" in events.columns else 0)

def _split_pipe_values(series: pd.Series) -> list[str]:
    values: list[str] = []
    if series is None:
        return values
    for raw in series.fillna("").astype(str):
        for item in raw.replace(";", "|").replace(",", "|").split("|"):
            item = re.sub(r"\s+", " ", item).strip()
            if item:
                values.append(item)
    return values


def _format_top_items(items: list[tuple[str, int]], limit: int = 5) -> str:
    filtered = [(str(name).strip(), int(count)) for name, count in items if str(name).strip()]
    if not filtered:
        return "нет выраженных лидеров"
    return "; ".join(f"{name} — {count}" for name, count in filtered[:limit])


def _summary_key(period_ids: list[str]) -> str:
    if period_ids:
        raw = "|".join(sorted(str(x) for x in period_ids if str(x).strip()))
        digest = hashlib.md5(raw.encode("utf-8", errors="ignore")).hexdigest()[:12]
        return f"periods:{digest}"
    return "local:current"




def _summary_period_range(messages, period_ids: list[str] | None = None) -> str:
    """Return human-readable period range for summary without period title.

    Prefer edited period metadata from Supabase. Message-level datetimes can be
    partially empty or parsed differently across imported sources, so using only
    messages may show the first selected period even when several periods are
    selected.
    """
    if period_ids:
        try:
            periods = list_periods(include_inactive=True)
            if isinstance(periods, pd.DataFrame) and not periods.empty and "period_id" in periods.columns:
                selected = periods[periods["period_id"].astype(str).isin([str(x) for x in period_ids])].copy()
                if not selected.empty:
                    dates = []
                    for col in ["date_from", "date_to"]:
                        if col in selected.columns:
                            parsed = pd.to_datetime(selected[col], errors="coerce", dayfirst=True).dropna()
                            if not parsed.empty:
                                dates.append(parsed.min())
                                dates.append(parsed.max())
                    if dates:
                        start_s = format_date(min(dates))
                        end_s = format_date(max(dates))
                        if start_s and end_s and start_s != end_s:
                            return f"{start_s} — {end_s}"
                        if start_s:
                            return start_s
        except Exception:
            pass

    if not isinstance(messages, pd.DataFrame) or messages.empty or "datetime" not in messages.columns:
        return "выбранный период"
    dt = pd.to_datetime(messages["datetime"], errors="coerce", dayfirst=True).dropna()
    if dt.empty:
        return "выбранный период"
    start_s = format_date(dt.min())
    end_s = format_date(dt.max())
    if start_s and end_s and start_s != end_s:
        return f"{start_s} — {end_s}"
    if start_s:
        return start_s
    return "выбранный период"


def _first_summary_line(summary_text: str) -> str:
    for line in str(summary_text or "").splitlines():
        cleaned = line.strip()
        if cleaned:
            return cleaned
    return ""


def _sync_summary_first_line(summary_text: str, auto_summary: str) -> str:
    """Keep manually saved summaries consistent with the current selected periods.

    Editors may save custom wording for the body of the summary, but the first
    line contains calculated metrics and selected period range. It should always
    be generated from the current data to avoid stale single-period headers when
    multiple periods are selected.
    """
    value = str(summary_text or "").strip()
    auto_first = _first_summary_line(auto_summary)
    if not value or not auto_first:
        return value or auto_summary
    lines = value.splitlines()
    for idx, line in enumerate(lines):
        if re.search(r"За период|За выбранный период", line):
            lines[idx] = auto_first
            return "\n".join(lines).strip()
    return "\n".join([auto_first, value]).strip()


def _events_grouped_by_title(events: pd.DataFrame, metric_col: str) -> pd.DataFrame:
    if not isinstance(events, pd.DataFrame) or events.empty or "event_title" not in events.columns:
        return pd.DataFrame(columns=["event_title", metric_col])
    work = events.copy()
    work["event_title"] = work["event_title"].fillna("").astype(str).str.strip()
    work = work[work["event_title"] != ""]
    if work.empty:
        return pd.DataFrame(columns=["event_title", metric_col])
    if metric_col not in work.columns:
        work[metric_col] = 0
    work[metric_col] = pd.to_numeric(work[metric_col], errors="coerce").fillna(0).astype(int)
    return work.groupby("event_title", as_index=False)[metric_col].sum().sort_values(metric_col, ascending=False)

def format_dashboard_summary_markdown(text: str) -> str:
    """Normalize summary display: remove legacy period title and bold leading labels."""
    import re

    value = str(text or "").strip()
    if not value:
        return ""

    value = re.sub(
        r"За период\s+«[^»]+»\s*\((\d{2}\.\d{2}\.\d{4}\s*[—-]\s*\d{2}\.\d{2}\.\d{4})\)",
        r"За период \1",
        value,
    )
    value = re.sub(
        r"За период\s+«[^»]+»\s*\((\d{2}\.\d{2}\.\d{4})\)",
        r"За период \1",
        value,
    )

    labels = [
        "Негатив",
        "Наиболее активные чаты",
        "Основные инфоповоды",
        "Чаще всего встречающиеся темы/теги",
        "Основные источники негатива",
    ]
    for label in labels:
        value = re.sub(rf"(^|\n)(\s*•\s*){re.escape(label)}:", rf"\1\2**{label}:**", value)
        value = re.sub(rf"(^|\n)(\s*){re.escape(label)}:", rf"\1\2**{label}:**", value)

    value = re.sub(
        r"(^|\n)(\s*•\s*)За период\s+([^\n]+?)\s+собрано",
        r"\1\2**За период \3** собрано",
        value,
    )
    value = re.sub(
        r"(^|\n)(\s*)За период\s+([^\n]+?)\s+собрано",
        r"\1\2**За период \3** собрано",
        value,
    )
    return value


def _period_label(period_ids: list[str]) -> str:
    if not period_ids:
        return "текущая выборка"
    try:
        periods = list_periods(include_inactive=True)
        if not periods.empty and "period_id" in periods.columns:
            selected = periods[periods["period_id"].astype(str).isin([str(x) for x in period_ids])].copy()
            if not selected.empty:
                names = selected.get("period_name", selected["period_id"]).fillna("").astype(str).tolist()
                names = [name for name in names if name]
                if len(names) == 1:
                    return names[0]
                if 1 < len(names) <= 3:
                    return ", ".join(names)
                if len(names) > 3:
                    return f"{len(names)} выбранных периода"
    except Exception:
        pass
    return f"{len(period_ids)} выбранных периода" if len(period_ids) != 1 else str(period_ids[0])


def build_auto_dashboard_summary(events: pd.DataFrame, messages: pd.DataFrame, period_ids: list[str]) -> str:
    """Build a readable editorial summary for the selected period(s)."""
    visible_messages = messages.copy() if isinstance(messages, pd.DataFrame) else pd.DataFrame()
    visible_events = events.copy() if isinstance(events, pd.DataFrame) else pd.DataFrame()

    if "message_hidden" in visible_messages.columns:
        visible_messages = visible_messages[~visible_messages["message_hidden"].astype(bool)]
    if "is_hidden" in visible_events.columns:
        visible_events = visible_events[~visible_events["is_hidden"].astype(bool)]

    period_range = _summary_period_range(visible_messages, period_ids)

    if visible_messages.empty and visible_events.empty:
        return "**За выбранный период** данных для саммари пока недостаточно."

    msg_count = int(len(visible_messages)) if not visible_messages.empty else int(visible_events.get("message_count", pd.Series(dtype=int)).sum())

    chat_col = "chat_title" if "chat_title" in visible_messages.columns else "chat_id" if "chat_id" in visible_messages.columns else None
    author_col = "author" if "author" in visible_messages.columns else "author_id" if "author_id" in visible_messages.columns else None
    chat_count = int(visible_messages[chat_col].nunique()) if chat_col and not visible_messages.empty else int(visible_events.get("chat_count", pd.Series(dtype=int)).max() if not visible_events.empty else 0)
    author_count = int(visible_messages[author_col].nunique()) if author_col and not visible_messages.empty else int(visible_events.get("author_count", pd.Series(dtype=int)).max() if not visible_events.empty else 0)

    if "is_negative" in visible_messages.columns and not visible_messages.empty:
        neg_mask = visible_messages["is_negative"].astype(str).str.lower().isin(["true", "1", "yes", "да"])
        negative_count = int(neg_mask.sum())
    elif "sentiment" in visible_messages.columns and not visible_messages.empty:
        neg_mask = visible_messages["sentiment"].astype(str).str.lower().str.contains("нег|neg", regex=True, na=False)
        negative_count = int(neg_mask.sum())
    else:
        negative_count = int(visible_events.get("negative_count", pd.Series(dtype=int)).sum() if not visible_events.empty else 0)
    negative_share = negative_count / msg_count if msg_count else 0.0

    top_chats_text = "нет данных"
    if chat_col and not visible_messages.empty:
        top_chats = (
            visible_messages[chat_col]
            .fillna("")
            .astype(str)
            .replace("", np.nan)
            .dropna()
            .value_counts()
            .head(5)
        )
        top_chats_text = _format_top_items(list(top_chats.items()))

    top_events_text = "нет выраженных тем"
    if not visible_events.empty and "event_title" in visible_events.columns:
        top_events = _events_grouped_by_title(visible_events, "message_count")
        if not top_events.empty:
            top_events_text = _format_top_items(list(zip(top_events["event_title"], top_events["message_count"])), limit=6)

    negative_events_text = "нет выраженного негативного ядра"
    if not visible_events.empty and {"event_title", "negative_count"}.issubset(visible_events.columns):
        neg_events = _events_grouped_by_title(visible_events, "negative_count")
        neg_events = neg_events[neg_events["negative_count"] > 0].head(5)
        if not neg_events.empty:
            negative_events_text = _format_top_items(list(zip(neg_events["event_title"], neg_events["negative_count"])))

    tag_values: list[str] = []
    if "main_tags" in visible_events.columns:
        tag_values.extend(_split_pipe_values(visible_events["main_tags"]))
    if "tags" in visible_messages.columns:
        tag_values.extend(_split_pipe_values(visible_messages["tags"]))
    tag_text = "нет явных тегов"
    if tag_values:
        top_tags = pd.Series(tag_values).value_counts().head(7)
        tag_text = _format_top_items(list(top_tags.items()), limit=7)

    lines = [
        f"**За период {period_range}** собрано {msg_count:,} сообщений из {chat_count:,} чатов".replace(",", " ") + (f" от {author_count:,} авторов".replace(",", " ") if author_count else "") + ".",
        f"**Негатив:** {negative_count:,} сообщений, доля — {format_pct(negative_share)}.".replace(",", " "),
        f"**Наиболее активные чаты:** {top_chats_text}.",
        f"**Основные инфоповоды:** {top_events_text}.",
        f"**Чаще всего встречающиеся темы/теги:** {tag_text}.",
        f"**Основные источники негатива:** {negative_events_text}.",
    ]
    return "\n".join(f"• {line}" for line in lines)


def render_dashboard_summary(events: pd.DataFrame, messages: pd.DataFrame, period_ids: list[str], conn, can_edit: bool) -> None:
    """Render editable period summary before the information-events block."""
    key = _summary_key(period_ids)
    auto_summary = build_auto_dashboard_summary(events, messages, period_ids)
    saved = get_dashboard_summary(conn, key)
    saved_text = str(saved.get("summary", "") or "").strip() if isinstance(saved, dict) else ""
    summary_text = _sync_summary_first_line(saved_text, auto_summary) if saved_text else auto_summary

    st.subheader("Саммари периода")
    if can_edit:
        if saved_text:
            st.caption("Показано ручное саммари. Автоматическое саммари можно вернуть, очистив ручную версию.")
        else:
            st.caption("Автоматическое саммари сформировано по выбранному периоду и текущей структуре инфоповодов.")

    st.markdown(format_dashboard_summary_markdown(summary_text).replace("\n", "  \n"))

    if not can_edit:
        return

    with st.expander("Редактировать саммари", expanded=False):
        with st.form(f"dashboard_summary_form_{key}"):
            edited_summary = st.text_area(
                "Саммари",
                value=summary_text,
                height=220,
                help="Этот текст будет показываться перед таблицей инфоповодов для выбранного периода или набора периодов.",
            )
            note = st.text_input("Комментарий к правке", value=str(saved.get("note", "") or "") if isinstance(saved, dict) else "")
            c1, c2 = st.columns(2)
            save_clicked = c1.form_submit_button("Сохранить саммари", type="primary")
            reset_clicked = c2.form_submit_button("Вернуть автоматическое")

        if save_clicked:
            save_dashboard_summary(conn, key, edited_summary.strip(), note=note, period_ids=period_ids)
            st.success("Саммари сохранено.")
            st.rerun()
        if reset_clicked:
            save_dashboard_summary(conn, key, "", note=note, period_ids=period_ids)
            st.success("Ручное саммари очищено. Будет показана автоматическая версия.")
            st.rerun()



# -----------------------------------------------------------------------------
# Digest export: DOCX / PDF
# -----------------------------------------------------------------------------


def _visible_messages_for_digest(messages: pd.DataFrame) -> pd.DataFrame:
    """Return messages that should be counted in the digest export."""
    if not isinstance(messages, pd.DataFrame) or messages.empty:
        return pd.DataFrame()
    work = messages.copy()
    if "message_hidden" in work.columns:
        work = work[~work["message_hidden"].astype(bool)]
    if "source_relevant" in work.columns:
        rel = work["source_relevant"].astype(str).str.lower().str.strip()
        work = work[~rel.isin(["false", "0", "no", "нет", "не релевантно", "нерелевантно"])]
    return work


def _visible_events_for_digest(events: pd.DataFrame) -> pd.DataFrame:
    if not isinstance(events, pd.DataFrame) or events.empty:
        return pd.DataFrame()
    work = events.copy()
    if "is_hidden" in work.columns:
        work = work[~work["is_hidden"].astype(bool)]
    return work


def _sentiment_counts(messages: pd.DataFrame) -> dict[str, int]:
    if not isinstance(messages, pd.DataFrame) or messages.empty:
        return {"neutral": 0, "negative": 0, "positive": 0, "total": 0}
    work = messages.copy()
    total = int(len(work))

    sentiment = work.get("sentiment", pd.Series([""] * len(work), index=work.index)).fillna("").astype(str).str.lower()
    negative = sentiment.str.contains("нег|neg|отриц", regex=True, na=False)
    positive = sentiment.str.contains("позит|pos|полож", regex=True, na=False)

    if "is_negative" in work.columns:
        negative = negative | work["is_negative"].astype(str).str.lower().isin(["true", "1", "yes", "да"])

    # If a system exports explicit neutral values, anything not positive/negative is neutral.
    neutral_count = max(0, total - int(negative.sum()) - int(positive.sum()))
    return {
        "neutral": int(neutral_count),
        "negative": int(negative.sum()),
        "positive": int(positive.sum()),
        "total": total,
    }


def _pct_of(count: int, total: int, decimals: int = 0) -> str:
    if not total:
        return "0%"
    value = count / total * 100
    if decimals:
        raw = f"{value:.{decimals}f}".replace(".", ",")
        raw = re.sub(r",0$", "", raw)
        return f"{raw}%"
    return f"{value:.0f}%"


def _clean_markdown_text(value: str) -> str:
    text = str(value or "").strip()
    text = text.replace("**", "")
    text = re.sub(r"^\s*•\s*", "", text)
    return text.strip()


def _summary_body_for_digest(events: pd.DataFrame, messages: pd.DataFrame, period_ids: list[str], conn) -> list[str]:
    """Return bullet-like digest lines without the calculated first period line."""
    key = _summary_key(period_ids)
    auto_summary = build_auto_dashboard_summary(events, messages, period_ids)
    saved = get_dashboard_summary(conn, key)
    saved_text = str(saved.get("summary", "") or "").strip() if isinstance(saved, dict) else ""
    summary_text = _sync_summary_first_line(saved_text, auto_summary) if saved_text else auto_summary
    lines = []
    for line in str(summary_text or "").splitlines():
        cleaned = _clean_markdown_text(line)
        if not cleaned:
            continue
        if cleaned.lower().startswith("за период") or cleaned.lower().startswith("за выбранный период"):
            continue
        lines.append(cleaned)
    if not lines:
        lines.append("В периоде выделены основные темы, динамика обсуждений и наиболее заметные источники негатива.")
    return lines


def _quote_rows_for_digest(
    topic_messages: pd.DataFrame,
    *,
    event_title: str,
    event_summary: str,
    tags: str,
    pinned_message_ids: set[str],
    limit: int,
) -> pd.DataFrame:
    if not isinstance(topic_messages, pd.DataFrame) or topic_messages.empty or limit <= 0:
        return pd.DataFrame()
    work = topic_messages.copy()
    if "message_id" in work.columns:
        work["message_id"] = work["message_id"].astype(str)
    else:
        work["message_id"] = ""

    pinned = work[work["message_id"].isin({str(x) for x in pinned_message_ids})].copy()
    if not pinned.empty and "datetime" in pinned.columns:
        pinned = pinned.sort_values("datetime")
    auto_limit = max(0, limit - len(pinned))
    auto_source = work[~work["message_id"].isin(set(pinned.get("message_id", [])))].copy()
    auto = rank_key_messages(
        auto_source,
        event_title=event_title,
        event_summary=event_summary,
        tags=tags,
        limit=auto_limit,
    ) if auto_limit else auto_source.iloc[0:0].copy()
    result = pd.concat([pinned, auto], ignore_index=False) if not pinned.empty else auto
    if result.empty:
        return result
    # Keep the digest readable and prevent repeated identical quotes.
    result = result.copy()
    result["_quote_key"] = result.get("text_clean", "").astype(str).apply(lambda x: _dedupe_key(x)[:120])
    result = result.drop_duplicates("_quote_key")
    return result.head(limit)


def build_digest_export_payload(
    events: pd.DataFrame,
    messages: pd.DataFrame,
    period_ids: list[str],
    conn,
    *,
    max_topics: int = 8,
    quotes_per_topic: int = 3,
) -> dict:
    """Build a plain data payload for DOCX/PDF digest export."""
    visible_messages = _visible_messages_for_digest(messages)
    visible_events = _visible_events_for_digest(events)
    period_range = _summary_period_range(visible_messages, period_ids)
    sentiment = _sentiment_counts(visible_messages)
    total_messages = int(sentiment["total"])

    pinned_df = get_key_message_pins(conn)
    if isinstance(pinned_df, pd.DataFrame) and not pinned_df.empty:
        pinned_df = pinned_df.copy()
        pinned_df["event_id"] = pinned_df.get("event_id", "").astype(str)
        pinned_df["message_id"] = pinned_df.get("message_id", "").astype(str)
    else:
        pinned_df = pd.DataFrame(columns=["event_id", "message_id"])

    summary_lines = _summary_body_for_digest(visible_events, visible_messages, period_ids, conn)

    topics: list[dict] = []
    if not visible_events.empty and "event_title" in visible_events.columns:
        work_events = visible_events.copy()
        work_events["event_title"] = work_events["event_title"].fillna("").astype(str).str.strip()
        work_events = work_events[work_events["event_title"] != ""]
        if "message_count" not in work_events.columns:
            work_events["message_count"] = 0
        work_events["message_count"] = pd.to_numeric(work_events["message_count"], errors="coerce").fillna(0).astype(int)

        grouped_rows = []
        for title, group in work_events.groupby("event_title", sort=False):
            event_ids = set(group.get("event_id", pd.Series(dtype=str)).astype(str).tolist())
            if "final_event_id" in visible_messages.columns:
                topic_messages = visible_messages[visible_messages["final_event_id"].astype(str).isin(event_ids)].copy()
            else:
                topic_messages = pd.DataFrame()
            msg_count = int(len(topic_messages)) if not topic_messages.empty else int(group["message_count"].sum())
            grouped_rows.append((title, group, event_ids, topic_messages, msg_count))

        grouped_rows = sorted(grouped_rows, key=lambda x: x[4], reverse=True)[:max_topics]

        for title, group, event_ids, topic_messages, msg_count in grouped_rows:
            if msg_count <= 0:
                continue
            topic_sentiment = _sentiment_counts(topic_messages) if not topic_messages.empty else {"neutral": 0, "negative": int(group.get("negative_count", pd.Series([0])).sum()), "positive": 0, "total": msg_count}
            # If only event-level negative is available, neutral is the rest.
            if topic_sentiment["total"] == 0:
                topic_sentiment["total"] = msg_count
            if topic_sentiment["neutral"] == 0 and topic_sentiment["positive"] == 0 and topic_sentiment["negative"] <= msg_count:
                topic_sentiment["neutral"] = max(0, msg_count - topic_sentiment["negative"])

            summaries = [str(x).strip() for x in group.get("event_summary", pd.Series(dtype=str)).fillna("").astype(str).tolist() if str(x).strip()]
            event_summary = summaries[0] if summaries else ""
            main_tags = "|".join(sorted(set(_split_pipe_values(group.get("main_tags", pd.Series(dtype=str)))))) if "main_tags" in group.columns else ""
            if not event_summary and not topic_messages.empty:
                event_summary = build_event_description(title, main_tags, topic_messages)

            pinned_message_ids = set(pinned_df.loc[pinned_df["event_id"].isin(event_ids), "message_id"].tolist())
            quote_rows = _quote_rows_for_digest(
                topic_messages,
                event_title=title,
                event_summary=event_summary,
                tags=main_tags,
                pinned_message_ids=pinned_message_ids,
                limit=quotes_per_topic,
            )
            quotes = []
            for _, row in quote_rows.iterrows():
                text = str(row.get("text_clean", "") or "").strip()
                if not text:
                    continue
                quotes.append({
                    "date": format_date(row.get("datetime")),
                    "chat": str(row.get("chat_title", "") or "").strip(),
                    "author": str(row.get("author", "") or "").strip(),
                    "text": re.sub(r"\s+", " ", text)[:900],
                    "link": str(row.get("message_link", "") or "").strip(),
                })

            topics.append({
                "title": title,
                "message_count": msg_count,
                "share": msg_count / total_messages if total_messages else 0.0,
                "sentiment": topic_sentiment,
                "summary": event_summary or "В теме обсуждались основные сообщения и реакции участников по выбранному инфоповоду.",
                "quotes": quotes,
            })

    return {
        "title": "Дайджест водительских чатов",
        "period_range": period_range,
        "message_count": total_messages,
        "sentiment": sentiment,
        "summary_lines": summary_lines,
        "topics": topics,
        "created_at": datetime.now().strftime("%d.%m.%Y %H:%M"),
    }


def _digest_filename(period_range: str, suffix: str) -> str:
    safe = re.sub(r"[^0-9A-Za-zА-Яа-я_.-]+", "_", str(period_range or "period"), flags=re.UNICODE).strip("_")
    safe = safe or "period"
    return f"digest_driver_chats_{safe}.{suffix}"


def _add_docx_bold_label(paragraph, label: str, text: str = ""):
    run = paragraph.add_run(label)
    run.bold = True
    if text:
        paragraph.add_run(text)


def generate_digest_docx(payload: dict) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as exc:
        raise RuntimeError("Для выгрузки Word установите зависимость python-docx.") from exc

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    title = f"{payload['title']} | {payload['period_range']}"
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(16)

    sentiment = payload["sentiment"]
    total = max(1, int(sentiment.get("total", 0)))
    doc.add_paragraph(f"Релевантных сообщений: {payload['message_count']:,}".replace(",", " "))
    doc.add_paragraph(
        "Тональность: "
        f"{_pct_of(sentiment.get('neutral', 0), total)} - нейтрал, "
        f"{_pct_of(sentiment.get('negative', 0), total)} - негатив, "
        f"{_pct_of(sentiment.get('positive', 0), total)} - позитив"
    )
    doc.add_paragraph("---")

    p = doc.add_paragraph()
    p.add_run("Главное за неделю").bold = True
    for line in payload.get("summary_lines", []):
        doc.add_paragraph(_clean_markdown_text(line), style=None).style = doc.styles["Normal"]
    doc.add_paragraph("---")

    p = doc.add_paragraph()
    p.add_run("Обсуждения недели").bold = True

    for idx, topic in enumerate(payload.get("topics", []), start=1):
        sent = topic["sentiment"]
        topic_total = max(1, int(sent.get("total", topic.get("message_count", 0))))
        header = (
            f"{topic['title']} — {_pct_of(topic.get('share', 0), 1, decimals=1)} сообщений | "
            f"Тональность: {_pct_of(sent.get('neutral', 0), topic_total)} нейтрал, "
            f"{_pct_of(sent.get('negative', 0), topic_total)} негатив, "
            f"{_pct_of(sent.get('positive', 0), topic_total)} позитив"
        )
        p = doc.add_paragraph()
        p.add_run(header).bold = True
        doc.add_paragraph(_clean_markdown_text(topic.get("summary", "")))
        quotes = topic.get("quotes", [])
        if quotes:
            qh = doc.add_paragraph()
            qh.add_run("Ключевые цитаты:").bold = True
            for quote in quotes:
                meta = " · ".join([x for x in [quote.get("date"), quote.get("chat"), quote.get("author")] if x])
                text = f"{meta}: {quote.get('text', '')}" if meta else quote.get("text", "")
                doc.add_paragraph(text, style="List Bullet")
        if idx != len(payload.get("topics", [])):
            doc.add_paragraph("---")

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _register_pdf_font():
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except Exception as exc:
        raise RuntimeError("Для выгрузки PDF установите зависимость reportlab.") from exc

    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            try:
                pdfmetrics.registerFont(TTFont("DigestSans", path))
                return "DigestSans"
            except Exception:
                continue
    # Fallback may not render Cyrillic in all environments, but keeps generation from crashing.
    return "Helvetica"


def generate_digest_pdf(payload: dict) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_LEFT
        from xml.sax.saxutils import escape as xml_escape
    except Exception as exc:
        raise RuntimeError("Для выгрузки PDF установите зависимость reportlab.") from exc

    font_name = _register_pdf_font()
    out = BytesIO()
    doc = SimpleDocTemplate(
        out,
        pagesize=A4,
        leftMargin=1.7 * cm,
        rightMargin=1.7 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
        title=f"{payload['title']} | {payload['period_range']}",
    )
    base = getSampleStyleSheet()
    normal = ParagraphStyle("DigestNormal", parent=base["Normal"], fontName=font_name, fontSize=9.8, leading=13, alignment=TA_LEFT)
    title_style = ParagraphStyle("DigestTitle", parent=normal, fontName=font_name, fontSize=16, leading=20, spaceAfter=10)
    h_style = ParagraphStyle("DigestHeading", parent=normal, fontName=font_name, fontSize=12.5, leading=16, spaceBefore=8, spaceAfter=6)
    topic_style = ParagraphStyle("DigestTopic", parent=normal, fontName=font_name, fontSize=10.5, leading=14, spaceBefore=8, spaceAfter=4)
    small = ParagraphStyle("DigestSmall", parent=normal, fontName=font_name, fontSize=8.8, leading=12, leftIndent=12)

    story = []
    story.append(Paragraph(f"<b>{xml_escape(payload['title'])} | {xml_escape(payload['period_range'])}</b>", title_style))
    sentiment = payload["sentiment"]
    total = max(1, int(sentiment.get("total", 0)))
    story.append(Paragraph(xml_escape(f"Релевантных сообщений: {payload['message_count']:,}".replace(",", " ")), normal))
    tonality = (
        "Тональность: "
        f"{_pct_of(sentiment.get('neutral', 0), total)} - нейтрал, "
        f"{_pct_of(sentiment.get('negative', 0), total)} - негатив, "
        f"{_pct_of(sentiment.get('positive', 0), total)} - позитив"
    )
    story.append(Paragraph(xml_escape(tonality), normal))
    story.append(Spacer(1, 8))
    story.append(Paragraph("<b>Главное за неделю</b>", h_style))
    for line in payload.get("summary_lines", []):
        story.append(Paragraph("• " + xml_escape(_clean_markdown_text(line)), normal))
    story.append(Spacer(1, 8))
    story.append(Paragraph("<b>Обсуждения недели</b>", h_style))

    for topic in payload.get("topics", []):
        sent = topic["sentiment"]
        topic_total = max(1, int(sent.get("total", topic.get("message_count", 0))))
        header = (
            f"{topic['title']} — {_pct_of(topic.get('share', 0), 1, decimals=1)} сообщений | "
            f"Тональность: {_pct_of(sent.get('neutral', 0), topic_total)} нейтрал, "
            f"{_pct_of(sent.get('negative', 0), topic_total)} негатив, "
            f"{_pct_of(sent.get('positive', 0), topic_total)} позитив"
        )
        story.append(Paragraph(f"<b>{xml_escape(header)}</b>", topic_style))
        story.append(Paragraph(xml_escape(_clean_markdown_text(topic.get("summary", ""))), normal))
        quotes = topic.get("quotes", [])
        if quotes:
            story.append(Paragraph("<b>Ключевые цитаты:</b>", small))
            for quote in quotes:
                meta = " · ".join([x for x in [quote.get("date"), quote.get("chat"), quote.get("author")] if x])
                text = f"{meta}: {quote.get('text', '')}" if meta else quote.get("text", "")
                story.append(Paragraph("• " + xml_escape(text), small))
        story.append(Spacer(1, 6))

    doc.build(story)
    return out.getvalue()


def render_digest_export(events: pd.DataFrame, messages: pd.DataFrame, period_ids: list[str], conn) -> None:
    """Render export controls for the current period selection."""
    with st.expander("Выгрузка дайджеста", expanded=False):
        st.caption("Сформируйте Word или PDF по выбранному периоду/периодам. В выгрузку попадают саммари, тональность, основные обсуждения и ключевые цитаты.")
        col1, col2 = st.columns(2)
        max_topics = int(col1.number_input("Тем в выгрузке", min_value=1, max_value=30, value=10, step=1))
        quotes_per_topic = int(col2.number_input("Цитат на тему", min_value=0, max_value=10, value=3, step=1))

        try:
            payload = build_digest_export_payload(
                events,
                messages,
                period_ids,
                conn,
                max_topics=max_topics,
                quotes_per_topic=quotes_per_topic,
            )
        except Exception as exc:
            st.error("Не удалось подготовить данные для выгрузки.")
            st.exception(exc)
            return

        if not payload.get("message_count"):
            st.info("Нет сообщений для выгрузки по выбранному периоду.")
            return

        st.write(
            f"Будет выгружено: {payload['message_count']:,} сообщений, "
            f"{len(payload.get('topics', []))} тем, период {payload.get('period_range', '')}.".replace(",", " ")
        )

        c1, c2 = st.columns(2)
        with c1:
            try:
                docx_bytes = generate_digest_docx(payload)
                st.download_button(
                    "Скачать Word",
                    data=docx_bytes,
                    file_name=_digest_filename(payload.get("period_range", "period"), "docx"),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as exc:
                st.warning(f"Word-выгрузка недоступна: {exc}")
        with c2:
            try:
                pdf_bytes = generate_digest_pdf(payload)
                st.download_button(
                    "Скачать PDF",
                    data=pdf_bytes,
                    file_name=_digest_filename(payload.get("period_range", "period"), "pdf"),
                    mime="application/pdf",
                    use_container_width=True,
                )
            except Exception as exc:
                st.warning(f"PDF-выгрузка недоступна: {exc}")


def _safe_int_delta(value) -> str:
    try:
        return f"{int(value):+d}"
    except Exception:
        return "0"



def _safe_pct_delta(current: float, previous: float) -> str:
    try:
        if previous in [0, None] or pd.isna(previous):
            return "н/д"
        return f"{((current - previous) / previous) * 100:+.0f}%"
    except Exception:
        return "н/д"



def build_period_comparison_df(messages: pd.DataFrame, period_ids: list[str], periods_meta: pd.DataFrame | None = None) -> pd.DataFrame:
    """Build period-level comparison metrics from visible messages."""
    if messages is None or messages.empty or "period_id" not in messages.columns:
        return pd.DataFrame()

    work = messages.copy()
    if period_ids:
        work = work[work["period_id"].astype(str).isin([str(x) for x in period_ids])]
    if work.empty:
        return pd.DataFrame()

    if "message_hidden" in work.columns:
        work = work[~work["message_hidden"].astype(bool)]
    if work.empty:
        return pd.DataFrame()

    if "datetime" in work.columns:
        work["datetime"] = pd.to_datetime(work["datetime"], errors="coerce")

    if "is_negative" in work.columns:
        work["__is_negative"] = work["is_negative"].astype(str).str.lower().isin(["true", "1", "yes", "да"]).astype(int)
    elif "sentiment" in work.columns:
        work["__is_negative"] = work["sentiment"].astype(str).str.lower().str.contains("neg").astype(int)
    else:
        work["__is_negative"] = 0

    agg_dict = {
        "message_count": ("period_id", "size"),
        "negative_count": ("__is_negative", "sum"),
    }
    if "datetime" in work.columns:
        agg_dict["date_from"] = ("datetime", "min")
        agg_dict["date_to"] = ("datetime", "max")
    if "chat_title" in work.columns:
        agg_dict["chat_count"] = ("chat_title", pd.Series.nunique)

    summary = work.groupby("period_id", dropna=False).agg(**agg_dict).reset_index()
    summary["negative_share"] = np.where(
        summary["message_count"] > 0,
        summary["negative_count"] / summary["message_count"],
        0.0,
    )

    if periods_meta is not None and not periods_meta.empty and "period_id" in periods_meta.columns:
        meta_cols = [c for c in ["period_id", "period_name", "date_from", "date_to", "uploaded_at"] if c in periods_meta.columns]
        meta = periods_meta[meta_cols].copy().drop_duplicates(subset=["period_id"])
        summary = summary.merge(meta, on="period_id", how="left", suffixes=("", "_meta"))
        if "date_from_meta" in summary.columns:
            summary["date_from"] = summary["date_from_meta"].combine_first(summary.get("date_from"))
        if "date_to_meta" in summary.columns:
            summary["date_to"] = summary["date_to_meta"].combine_first(summary.get("date_to"))
        drop_cols = [c for c in ["date_from_meta", "date_to_meta"] if c in summary.columns]
        if drop_cols:
            summary = summary.drop(columns=drop_cols)
    else:
        summary["period_name"] = summary["period_id"]

    # Normalize period names to plain strings. Supabase/JSON values can occasionally
    # come back as mixed objects; sorting mixed Python objects may crash pandas.
    if "period_name" not in summary.columns:
        summary["period_name"] = summary["period_id"]
    summary["period_name"] = summary["period_name"].where(
        summary["period_name"].notna(),
        summary["period_id"],
    ).map(lambda value: "" if pd.isna(value) else str(value))

    # Build a stable string sort key instead of sorting raw datetimes.
    # This avoids pandas TypeError when selected periods contain a mix of
    # timezone-aware, timezone-naive, empty, or malformed dates.
    date_from_series = summary["date_from"] if "date_from" in summary.columns else pd.Series([pd.NaT] * len(summary), index=summary.index)
    sort_date = pd.to_datetime(date_from_series, errors="coerce", utc=True)
    if "uploaded_at" in summary.columns:
        uploaded_at = pd.to_datetime(summary["uploaded_at"], errors="coerce", utc=True)
        sort_date = sort_date.combine_first(uploaded_at)

    summary["sort_date"] = sort_date.dt.strftime("%Y-%m-%d %H:%M:%S").fillna("9999-12-31 23:59:59")
    summary["period_name_sort"] = summary["period_name"].fillna("").astype(str)
    summary = summary.sort_values(
        ["sort_date", "period_name_sort"],
        ascending=[True, True],
        kind="mergesort",
    ).reset_index(drop=True)
    summary["negative_share_pct"] = (summary["negative_share"] * 100).round(1)
    summary["label"] = summary["period_name"]
    return summary



def render_period_comparison(messages: pd.DataFrame, period_ids: list[str]):
    """Render visual comparison across two or more selected periods."""
    if not period_ids or len(period_ids) < 2:
        return

    try:
        periods_meta = list_periods()
    except Exception:
        periods_meta = pd.DataFrame()

    summary = build_period_comparison_df(messages, period_ids, periods_meta)
    if summary.empty or len(summary) < 2:
        return

    latest = summary.iloc[-1]
    previous = summary.iloc[-2]
    msg_delta = int(latest["message_count"] - previous["message_count"])
    neg_count_delta = int(latest["negative_count"] - previous["negative_count"])
    neg_share_delta_pp = float((latest["negative_share"] - previous["negative_share"]) * 100)

    st.subheader("Сравнение периодов")
    st.caption(
        f"Сейчас сравниваются {len(summary)} период(а/ов). Последний период: {latest['period_name']}. "
        f"Сравнение ведется с предыдущим периодом: {previous['period_name']}."
    )

    c1, c2, c3, c4 = st.columns(4)
    c1.metric(
        f"Сообщения · {latest['period_name']}",
        f"{int(latest['message_count']):,}".replace(",", " "),
        delta=_safe_int_delta(msg_delta),
        help="Разница по общему количеству сообщений относительно предыдущего выбранного периода.",
    )
    c2.metric(
        "Изменение сообщений",
        _safe_pct_delta(float(latest["message_count"]), float(previous["message_count"])),
        help="Процентное изменение количества сообщений относительно предыдущего периода.",
    )
    c3.metric(
        f"Негатив · {latest['period_name']}",
        format_pct(latest["negative_share"]),
        delta=f"{neg_share_delta_pp:+.1f} п.п.",
        help="Доля негативных сообщений и ее изменение в процентных пунктах относительно предыдущего периода.",
    )
    c4.metric(
        "Негативных сообщений",
        f"{int(latest['negative_count']):,}".replace(",", " "),
        delta=_safe_int_delta(neg_count_delta),
        help="Изменение числа негативных сообщений относительно предыдущего периода.",
    )

    chart_source = summary[["label", "message_count", "negative_share_pct"]].copy().rename(columns={
        "label": "Период",
        "message_count": "Сообщений",
        "negative_share_pct": "Негатив, %",
    })

    left, right = st.columns(2)
    with left:
        st.caption("Общее количество сообщений по периодам")
        st.bar_chart(chart_source.set_index("Период")[["Сообщений"]], use_container_width=True)
    with right:
        st.caption("Доля негатива по периодам")
        st.line_chart(chart_source.set_index("Период")[["Негатив, %"]], use_container_width=True)

    table = summary.copy()
    table["Период"] = table["period_name"]
    table["Начало"] = format_date_series(table.get("date_from", pd.Series(dtype=str)))
    table["Конец"] = format_date_series(table.get("date_to", pd.Series(dtype=str)))
    table["Сообщений"] = table["message_count"].astype(int)
    table["Негативных"] = table["negative_count"].astype(int)
    table["Доля негатива"] = table["negative_share"].apply(format_pct)
    display_cols = [c for c in ["Период", "Начало", "Конец", "Сообщений", "Негативных", "Доля негатива"] if c in table.columns]
    st.dataframe(table[display_cols], use_container_width=True, hide_index=True)



def word_message_results_table(messages: pd.DataFrame, events: pd.DataFrame, query: str):
    """Show all messages that match the sidebar word filter."""
    if not str(query or "").strip():
        return

    st.subheader(f"Сообщения со словом: «{query}»")

    if messages.empty:
        st.info("Сообщений с таким словом в тексте не найдено.")
        return

    table = messages.copy().sort_values("datetime", ascending=False)
    table["Дата"] = format_date_series(table.get("datetime", pd.Series(dtype=str)))
    table["Чат"] = table.get("chat_title", "")
    table["Автор"] = table.get("author", "")
    table["Текст"] = table.get("text_clean", "").fillna("").astype(str).str.slice(0, 500)
    event_title_map = events.set_index("event_id")["event_title"].to_dict() if len(events) else {}
    table["Инфоповод"] = table.get("final_event_id", "").map(event_title_map).fillna("")
    table["Ссылка"] = table.get("message_link", "")

    st.caption(f"Найдено сообщений: {len(table):,}".replace(",", " "))
    cols = ["Дата", "Чат", "Автор", "Инфоповод", "Текст", "Ссылка"]
    cols = [c for c in cols if c in table.columns]

    st.dataframe(
        table[cols],
        use_container_width=True,
        height=420,
        hide_index=True,
        column_config={
            "Текст": st.column_config.TextColumn(width="large"),
            "Инфоповод": st.column_config.TextColumn(width="medium"),
            "Ссылка": st.column_config.LinkColumn("Ссылка"),
        },
    )


def event_table(events: pd.DataFrame) -> str | None:
    if events.empty:
        st.info("Нет инфоповодов по выбранным фильтрам.")
        return None

    table = events.copy()
    table["Период"] = table.apply(format_period, axis=1)
    table["Негатив"] = table["negative_share"].apply(format_pct)
    table["Токсичность"] = table["toxic_share"].apply(format_pct)
    table["Теги"] = table["main_tags"].fillna("").astype(str).str.replace("|", ", ", regex=False)
    table["Тема из файла"] = table.get("source_main_topics", pd.Series([""] * len(table))).fillna("").astype(str).str.replace(";", ",", regex=False)
    table["Название"] = table["event_title"]
    table["Описание"] = table["event_summary"]
    table["Сообщений"] = table["message_count"]
    table["Чатов"] = table["chat_count"]
    table["Авторов"] = table["author_count"]
    table["Важность"] = table["importance_score"].round(1)
    table["Статус"] = table["status"]

    display_cols = [
        "Название",
        "Описание",
    ]
    if table["Тема из файла"].fillna("").astype(str).str.strip().any():
        display_cols.append("Тема из файла")
    display_cols += [
        "Теги",
        "Период",
        "Сообщений",
        "Чатов",
        "Негатив",
        "Важность",
        "Статус",
    ]

    st.subheader("Инфоповоды")
    selected = st.dataframe(
        table[display_cols],
        use_container_width=True,
        height=520,
        hide_index=True,
        on_select="rerun",
        selection_mode="single-row",
        column_config={
            "Описание": st.column_config.TextColumn(width="large"),
            "Название": st.column_config.TextColumn(width="medium"),
            "Теги": st.column_config.TextColumn(width="medium"),
            "Тема из файла": st.column_config.TextColumn(width="medium"),
            "Сообщений": st.column_config.NumberColumn(format="%d"),
            "Чатов": st.column_config.NumberColumn(format="%d"),
            "Важность": st.column_config.NumberColumn(format="%.1f"),
        },
    )

    rows = get_selected_rows(selected)
    if rows:
        return table.iloc[rows[0]]["event_id"]
    return table.iloc[0]["event_id"]


KEY_MESSAGE_STOPWORDS = {
    "это", "как", "что", "или", "для", "при", "уже", "еще", "ещё", "там", "тут", "все", "всё",
    "они", "она", "оно", "его", "ему", "вам", "вас", "нам", "нас", "про", "без", "под", "над",
    "только", "так", "вот", "когда", "если", "где", "куда", "почему", "зачем", "тоже", "можно",
    "нужно", "надо", "будет", "было", "были", "есть", "нет", "да", "ну", "же", "ли", "бы", "за",
    "из", "от", "до", "по", "на", "не", "ни", "во", "со", "ко", "то", "вы", "мы", "он", "я",
    "такси", "чат", "водитель", "водители", "сообщение", "сообщения", "тема", "обсуждение",
}


def _message_tokens(text: str) -> list[str]:
    """Tokenize Russian/English text for representative-message scoring."""
    normalized = str(text or "").lower().replace("ё", "е")
    tokens = re.findall(r"[a-zа-я0-9]{3,}", normalized, flags=re.IGNORECASE)
    return [t for t in tokens if t not in KEY_MESSAGE_STOPWORDS and not t.isdigit()]


def _message_token_set(text: str) -> set[str]:
    return set(_message_tokens(text))


def _dedupe_key(text: str) -> str:
    normalized = re.sub(r"\s+", " ", str(text or "").lower().replace("ё", "е")).strip()
    normalized = re.sub(r"[^a-zа-я0-9 ]+", "", normalized)
    return normalized[:280]


def _text_source_penalty(row: pd.Series) -> float:
    source = str(row.get("text_source", "") or row.get("source", "")).lower()
    text = str(row.get("text_clean", "") or "").lower()
    penalty = 0.0
    if any(x in source for x in ["ocr", "image", "изображ", "картин"]):
        penalty -= 0.35
    if any(x in text for x in ["тексты с изображений", "расшифровки"]):
        penalty -= 0.25
    return penalty


def rank_key_messages(
    event_messages: pd.DataFrame,
    *,
    event_title: str = "",
    event_summary: str = "",
    tags: str = "",
    limit: int = 10,
) -> pd.DataFrame:
    """Return representative messages for an event.

    The previous version prioritized negative + longest messages. This scorer keeps that signal,
    but adds semantic relevance to the event title/summary, frequent words inside the event,
    moderate text length, duplicate suppression and simple diversity.
    """
    if event_messages is None or event_messages.empty or "text_clean" not in event_messages.columns:
        return event_messages.iloc[0:0].copy() if isinstance(event_messages, pd.DataFrame) else pd.DataFrame()

    work = event_messages.copy()
    work["_text"] = work["text_clean"].fillna("").astype(str).str.strip()
    work = work[work["_text"].str.len() >= 18].copy()
    if work.empty:
        return event_messages.copy().head(limit)

    topic_seed = " ".join([str(event_title or ""), str(event_summary or ""), str(tags or "").replace("|", " ")])
    topic_tokens = set(_message_tokens(topic_seed))

    all_tokens: list[str] = []
    for txt in work["_text"].head(400).tolist():
        all_tokens.extend(_message_tokens(txt))
    frequent_tokens = {token for token, count in Counter(all_tokens).most_common(24) if count >= 2}
    relevance_tokens = topic_tokens | frequent_tokens

    rows = []
    seen_best: dict[str, tuple[float, int]] = {}
    for idx, row in work.iterrows():
        text = row.get("_text", "")
        text_len = len(text)
        tokens = _message_token_set(text)
        if not tokens:
            continue

        if text_len < 45:
            length_score = 0.20
        elif text_len <= 450:
            length_score = 0.85 + min(text_len, 450) / 450 * 0.35
        elif text_len <= 1200:
            length_score = 1.05
        else:
            length_score = 0.70

        overlap_topic = len(tokens & topic_tokens)
        overlap_relevance = len(tokens & relevance_tokens)
        topic_score = min(overlap_topic * 0.85, 3.4)
        relevance_score = min(overlap_relevance * 0.28, 2.2)

        sentiment = str(row.get("sentiment", "") or "").lower()
        is_negative = False
        if "is_negative" in row.index:
            is_negative = str(row.get("is_negative", "")).lower() in {"true", "1", "yes", "да"}
        is_negative = is_negative or any(x in sentiment for x in ["негатив", "negative", "отриц"])
        negative_score = 1.15 if is_negative else 0.0

        toxicity = str(row.get("toxicity", "") or row.get("Токсичность", "")).lower()
        toxic_score = 0.45 if any(x in toxicity for x in ["токс", "toxic", "да", "true", "1"]) else 0.0

        problem_score = 0.0
        problem_words = {
            "сбой", "ошибка", "ошибки", "проблема", "проблемы", "не работает",
            "заблокировали", "блокировка", "налог", "закон", "коэффициент", "приоритет",
            "оплата", "выплата", "штраф", "забастовка", "обновление", "тариф",
        }
        text_l = text.lower().replace("ё", "е")
        for word in problem_words:
            if word in text_l:
                problem_score += 0.22
        problem_score = min(problem_score, 1.1)

        score = length_score + topic_score + relevance_score + negative_score + toxic_score + problem_score + _text_source_penalty(row)

        if re.fullmatch(r"[а-яa-z0-9 ,.!?\-]{0,80}", text_l) and len(tokens) <= 4 and overlap_relevance == 0:
            score -= 1.0

        dedupe = _dedupe_key(text)
        if not dedupe:
            continue
        if dedupe in seen_best:
            old_score, old_idx = seen_best[dedupe]
            if score > old_score:
                seen_best[dedupe] = (score, idx)
        else:
            seen_best[dedupe] = (score, idx)
        rows.append((idx, score, tokens))

    if not rows:
        return work.sort_values("datetime").head(limit) if "datetime" in work.columns else work.head(limit)

    score_by_idx = {idx: score for idx, score, _ in rows}
    token_by_idx = {idx: tokens for idx, _, tokens in rows}
    duplicate_winners = {idx for _, idx in seen_best.values()}

    candidates = [idx for idx in score_by_idx if idx in duplicate_winners]
    candidates = sorted(candidates, key=lambda idx: score_by_idx[idx], reverse=True)

    selected: list[int] = []
    for idx in candidates:
        tokens = token_by_idx.get(idx, set())
        too_similar = False
        for chosen in selected:
            chosen_tokens = token_by_idx.get(chosen, set())
            union = tokens | chosen_tokens
            if union and len(tokens & chosen_tokens) / len(union) >= 0.72:
                too_similar = True
                break
        if too_similar:
            continue
        selected.append(idx)
        if len(selected) >= limit:
            break

    if len(selected) < limit:
        for idx in candidates:
            if idx not in selected:
                selected.append(idx)
            if len(selected) >= limit:
                break

    result = work.loc[selected].copy()
    result["_key_message_score"] = result.index.map(score_by_idx).fillna(0.0)
    if "datetime" in result.columns:
        return result.sort_values("datetime")
    return result


def message_preview_cards(
    event_messages: pd.DataFrame,
    limit: int = 8,
    *,
    event_title: str = "",
    event_summary: str = "",
    tags: str = "",
    pinned_message_ids: set[str] | None = None,
    conn=None,
    event_id: str = "",
    can_edit: bool = False,
):
    if event_messages.empty:
        st.info("Сообщения не найдены.")
        return

    pinned_message_ids = {str(x) for x in (pinned_message_ids or set()) if str(x).strip()}
    work = event_messages.copy()
    work["message_id"] = work.get("message_id", "").astype(str)

    pinned_sample = work[work["message_id"].isin(pinned_message_ids)].copy()
    auto_source = work[~work["message_id"].isin(pinned_message_ids)].copy()

    auto_limit = max(0, limit - len(pinned_sample))
    auto_sample = rank_key_messages(
        auto_source,
        event_title=event_title,
        event_summary=event_summary,
        tags=tags,
        limit=auto_limit,
    ) if auto_limit else auto_source.iloc[0:0].copy()

    if not pinned_sample.empty:
        pinned_sample["_manual_key"] = True
        if "datetime" in pinned_sample.columns:
            pinned_sample = pinned_sample.sort_values("datetime")
    if not auto_sample.empty:
        auto_sample["_manual_key"] = False

    sample = pd.concat([pinned_sample, auto_sample], ignore_index=False) if not pinned_sample.empty else auto_sample

    if sample.empty:
        st.info("Не найдено достаточно информативных сообщений для этого инфоповода.")
        return

    st.caption(
        "Закрепленные вручную сообщения показываются первыми. Остальные выбраны по близости к теме, частым словам внутри инфоповода, информативности текста, негативу и отсутствию дублей."
    )

    for _, row in sample.iterrows():
        when = format_date(row.get("datetime"))
        chat = row.get("chat_title", "")
        author = row.get("author", "")
        text = str(row.get("text_clean", "")).strip()
        sentiment = str(row.get("sentiment", "")).strip()
        link = str(row.get("message_link", "")).strip()
        message_id = str(row.get("message_id", "")).strip()
        is_manual_key = bool(row.get("_manual_key", False))
        badge = "<b>Закреплено вручную</b> · " if is_manual_key else ""

        st.markdown(
            f"""
<div style="padding: 0.75rem 0; border-bottom: 1px solid rgba(128,128,128,.25);">
  <div style="font-size: 0.88rem; opacity: .75;">{badge}{when} · {chat} · {author} · {sentiment}</div>
  <div style="margin-top: .25rem; white-space: pre-wrap;">{text[:1200]}</div>
</div>
""",
            unsafe_allow_html=True,
        )
        cols = st.columns([1, 4]) if can_edit and is_manual_key and conn is not None and event_id and message_id else None
        if link.startswith("http"):
            st.markdown(f"[Открыть сообщение]({link})")
        if cols:
            with cols[0]:
                if st.button("Убрать из ключевых", key=f"unpin_key_{event_id}_{message_id}"):
                    unpin_key_message(conn, event_id, message_id)
                    st.success("Сообщение удалено из ключевых.")
                    st.cache_data.clear()
                    st.rerun()


def show_event_card(event_id: str, events: pd.DataFrame, messages: pd.DataFrame, conn, can_edit: bool = True):
    selected = events[events["event_id"] == event_id]
    if selected.empty:
        st.warning("Инфоповод не найден.")
        return

    ev = selected.iloc[0]
    event_messages = messages[
        (messages["final_event_id"] == event_id)
        & (~messages.get("message_hidden", pd.Series([False] * len(messages))).astype(bool))
    ].copy()
    event_messages = event_messages.sort_values("datetime")

    pinned_messages = get_key_message_pins(conn)
    if pinned_messages is not None and not pinned_messages.empty:
        pinned_messages = pinned_messages.copy()
        pinned_messages["event_id"] = pinned_messages.get("event_id", "").astype(str)
        pinned_messages["message_id"] = pinned_messages.get("message_id", "").astype(str)
        pinned_message_ids = set(pinned_messages.loc[pinned_messages["event_id"] == str(event_id), "message_id"].tolist())
    else:
        pinned_message_ids = set()

    st.markdown("---")
    st.header(ev["event_title"])
    st.info(str(ev.get("event_summary", "")))

    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Сообщений", int(ev.get("message_count", 0)))
    c2.metric("Чатов", int(ev.get("chat_count", 0)))
    c3.metric("Авторов", int(ev.get("author_count", 0)))
    c4.metric("Негатив", format_pct(ev.get("negative_share", 0)))
    c5.metric("Важность", round(float(ev.get("importance_score", 0)), 1))

    tags = str(ev.get("main_tags", "")).replace("|", " · ")
    source_topics_caption = str(ev.get("source_main_topics", "") or "").strip()

    if source_topics_caption:
        st.caption(f"Тема из файла: {source_topics_caption}")
    st.caption(f"Теги: {tags}")

    tab_names = ["Ключевые сообщения", "Вся лента"] + (["Правки"] if can_edit else [])
    tabs = st.tabs(tab_names)
    tab_messages = tabs[0]
    tab_all = tabs[1]
    tab_edit = tabs[2] if can_edit else None

    with tab_messages:
        message_preview_cards(
            event_messages,
            limit=10,
            event_title=str(ev.get("event_title", "")),
            event_summary=str(ev.get("event_summary", "")),
            tags=tags,
            pinned_message_ids=pinned_message_ids,
            conn=conn,
            event_id=event_id,
            can_edit=can_edit,
        )

    with tab_all:
        if event_messages.empty:
            st.info("Сообщения не найдены.")
        else:
            table = event_messages.copy()
            table["Дата"] = format_date_series(table["datetime"])
            table["Текст"] = table["text_clean"].fillna("").astype(str).str.slice(0, 420)
            table["Тема"] = table.apply(lambda r: message_topic_display(r, str(ev.get("event_title", ""))), axis=1)
            table["Теги"] = table.get("tags", "").apply(format_tags_for_display) if "tags" in table.columns else ""
            table["Чат"] = table.get("chat_title", "")
            table["Автор"] = table.get("author", "")
            table["Тональность"] = table.get("sentiment", "")
            table["Ссылка"] = table.get("message_link", "")

            cols = ["Дата", "Текст", "Тема", "Теги", "Чат", "Автор", "Тональность", "Ссылка"]
            msg_select = st.dataframe(
                table[cols],
                use_container_width=True,
                height=460,
                hide_index=True,
                on_select="rerun",
                selection_mode="single-row",
                column_config={
                    "Дата": st.column_config.TextColumn("Дата", width="small"),
                    "Текст": st.column_config.TextColumn("Текст", width="large"),
                    "Тема": st.column_config.TextColumn("Тема", width="medium"),
                    "Теги": st.column_config.TextColumn("Теги", width="medium"),
                    "Ссылка": st.column_config.LinkColumn("Ссылка", width="small"),
                },
            )

            rows = get_selected_rows(msg_select)
            if rows:
                row = table.iloc[rows[0]]
                message_id = str(row.get("message_id", "")).strip()
                st.markdown("#### Выбранное сообщение")
                st.caption(
                    f"{format_date(row.get('datetime'))} · {row.get('chat_title', '')} · "
                    f"{row.get('author', '')} · текущая тема: {message_topic_display(row, str(ev.get('event_title', '')))}"
                )
                st.write(row.get("text_clean", ""))

                if can_edit:
                    with st.expander("Назначить тему или инфоповод", expanded=True):
                        msg_note = st.text_input(
                            "Комментарий к действию",
                            value="",
                            key=f"msg_note_{event_id}_{message_id}",
                        )

                        st.markdown("##### Назначить тему")
                        available_topics = collect_available_source_topics(events, messages)
                        current_topic = message_topic_display(row, str(ev.get("event_title", "")))
                        topic_options = [""] + available_topics
                        default_topic_idx = 0
                        for i, item in enumerate(topic_options):
                            if item and item.lower().replace("ё", "е") == current_topic.lower().replace("ё", "е"):
                                default_topic_idx = i
                                break
                        selected_topic = st.selectbox(
                            "Тема сообщения",
                            options=topic_options,
                            index=default_topic_idx,
                            format_func=lambda x: x or "— выбрать из найденных тем —",
                            key=f"topic_select_{event_id}_{message_id}",
                        )
                        custom_topic = st.text_input(
                            "Или новая тема",
                            value="",
                            placeholder="Например: Детские кресла",
                            key=f"custom_topic_{event_id}_{message_id}",
                        )
                        topic_to_save = custom_topic.strip() or selected_topic.strip()
                        if st.button(
                            "Сохранить тему сообщения",
                            disabled=not bool(topic_to_save),
                            key=f"save_topic_{event_id}_{message_id}",
                        ):
                            try:
                                save_message_topic_override(
                                    conn,
                                    message_id,
                                    source_main_topic=topic_to_save,
                                    source_topics=topic_to_save,
                                    note=msg_note,
                                )
                                st.success("Тема сообщения сохранена.")
                                st.cache_data.clear()
                                st.rerun()
                            except Exception as e:
                                st.error(str(e))

                        st.markdown("##### Назначить инфоповод")
                        target_options = events[["event_id", "event_title", "message_count", "start_date", "end_date"]].copy()
                        target_options["event_id"] = target_options["event_id"].astype(str)
                        target_options = target_options.sort_values(["message_count", "event_title"], ascending=[False, True])
                        event_search = st.text_input(
                            "Поиск инфоповода",
                            value="",
                            placeholder="Начните вводить название темы",
                            key=f"event_search_{event_id}_{message_id}",
                        )
                        if event_search.strip():
                            q = event_search.strip().lower().replace("ё", "е")
                            target_options = target_options[
                                target_options["event_title"].fillna("").astype(str).str.lower().str.replace("ё", "е", regex=False).str.contains(q, regex=False, na=False)
                            ]

                        target_label_map = {}
                        for _, target_row in target_options.iterrows():
                            target_id = str(target_row.get("event_id", ""))
                            title_part = str(target_row.get("event_title", "") or "Без названия")[:120]
                            msg_count = int(target_row.get("message_count", 0) or 0)
                            period_part = format_period(target_row)
                            suffix = f" · {msg_count} сообщ."
                            if period_part:
                                suffix += f" · {period_part}"
                            target_label_map[target_id] = f"{title_part}{suffix}"

                        target_ids = list(target_label_map.keys())
                        current_target_idx = 0
                        if str(event_id) in target_ids:
                            current_target_idx = target_ids.index(str(event_id))
                        selected_target_id = st.selectbox(
                            "Инфоповод сообщения",
                            options=target_ids,
                            index=current_target_idx if target_ids else None,
                            format_func=lambda x: target_label_map.get(str(x), str(x)),
                            key=f"target_event_id_{event_id}_{message_id}",
                            disabled=not bool(target_ids),
                        )
                        if selected_target_id:
                            st.caption(f"Будет назначен инфоповод: {target_label_map.get(str(selected_target_id), selected_target_id)}")

                        action_cols = st.columns(4)
                        if action_cols[0].button(
                            "Назначить инфоповод",
                            disabled=not bool(selected_target_id) or str(selected_target_id) == str(event_id),
                            key=f"assign_event_{event_id}_{message_id}",
                        ):
                            try:
                                move_message(conn, message_id, str(selected_target_id), note=msg_note)
                                st.success("Сообщению назначен выбранный инфоповод.")
                                st.cache_data.clear()
                                st.rerun()
                            except Exception as e:
                                st.error(str(e))

                        if action_cols[1].button("Нерелевант", key=f"irrelevant_{event_id}_{message_id}"):
                            mark_message_irrelevant(conn, event_id, message_id, reason=msg_note)
                            st.success("Сообщение исключено из текущего инфоповода как нерелевантное.")
                            st.cache_data.clear()
                            st.rerun()

                        if action_cols[2].button("Скрыть везде", key=f"hide_{event_id}_{message_id}"):
                            hide_message(conn, message_id, hidden=True, note=msg_note)
                            st.success("Сообщение скрыто во всех разделах дашборда.")
                            st.cache_data.clear()
                            st.rerun()

                        already_key = message_id in pinned_message_ids
                        if action_cols[3].button("Ключевое", disabled=already_key, key=f"pin_key_{event_id}_{message_id}"):
                            pin_key_message(conn, event_id, message_id, note=msg_note)
                            st.success("Сообщение добавлено в ключевые.")
                            st.cache_data.clear()
                            st.rerun()
                        if already_key:
                            st.caption("Это сообщение уже закреплено как ключевое для текущего инфоповода.")

                        st.markdown("##### Создать новый инфоповод и назначить его сообщению")
                        st.caption("Если подходящего инфоповода нет в списке, создайте новый — выбранное сообщение сразу будет перенесено туда.")
                        with st.form(f"create_event_for_message_{event_id}_{message_id}"):
                            new_title = st.text_input("Название нового инфоповода", key=f"new_event_title_{event_id}_{message_id}")
                            new_summary = st.text_area(
                                "Описание нового инфоповода",
                                value="",
                                height=90,
                                key=f"new_event_summary_{event_id}_{message_id}",
                            )
                            new_tags = st.text_input(
                                "Теги нового инфоповода",
                                value=str(row.get("tags", "")).replace("|", ", "),
                                key=f"new_event_tags_{event_id}_{message_id}",
                            )
                            create_and_move = st.form_submit_button("Создать и назначить")
                            if create_and_move:
                                try:
                                    new_event_id = create_manual_event(
                                        conn,
                                        title=new_title,
                                        summary=new_summary,
                                        status="новый",
                                        main_tags=normalize_manual_tags(new_tags),
                                        note=f"Создано из сообщения {message_id}",
                                    )
                                    move_message(conn, message_id, new_event_id, note="Назначено в новый инфоповод")
                                    st.success("Новый инфоповод создан, сообщение назначено в него.")
                                    st.cache_data.clear()
                                    st.rerun()
                                except Exception as e:
                                    st.error(str(e))

    if can_edit and tab_edit is not None:
        with tab_edit:
            st.markdown("#### Ручная правка инфоповода")
            st.caption("Описание можно скорректировать вручную. После сохранения оно будет показываться в таблице и карточке вместо автоматического описания.")
            with st.form(f"edit_event_{event_id}"):
                title = st.text_input("Название", value=str(ev.get("event_title", "")))
                summary = st.text_area("Описание", value=str(ev.get("event_summary", "")), height=150, help="Например: В теме обсуждались: законопроект; повышение коэффициентов; невозможность загрузить обновление.")
                status = st.selectbox(
                    "Статус",
                    STATUS_OPTIONS,
                    index=STATUS_OPTIONS.index(ev.get("status")) if ev.get("status") in STATUS_OPTIONS else 0,
                )
                hidden = st.checkbox("Скрыть инфоповод", value=bool(ev.get("is_hidden", False)))
                note = st.text_area("Комментарий модератора", value="", height=80)
                submitted = st.form_submit_button("Сохранить")
                if submitted:
                    save_event_override(conn, event_id, title=title, summary=summary, status=status, hidden=hidden, note=note)
                    st.success("Правки сохранены.")
                    st.cache_data.clear()
                    st.rerun()

            st.markdown("#### Нерелевантные сообщения")
            exclusions = get_message_exclusions(conn)
            event_exclusions = exclusions[exclusions["event_id"] == event_id] if not exclusions.empty else exclusions
            if event_exclusions.empty:
                st.info("Для этого инфоповода пока нет сообщений, помеченных как нерелевантные.")
            else:
                excluded_ids = event_exclusions["message_id"].astype(str).tolist()
                excluded_messages = messages[messages["message_id"].astype(str).isin(excluded_ids)].copy()
                if excluded_messages.empty:
                    st.info("Есть записи об исключениях, но сообщения не найдены в текущей выгрузке.")
                else:
                    reason_map = event_exclusions.set_index("message_id")["reason"].to_dict()
                    excluded_messages["Дата"] = format_date_series(excluded_messages.get("datetime", pd.Series(dtype=str)))
                    excluded_messages["Чат"] = excluded_messages.get("chat_title", "")
                    excluded_messages["Автор"] = excluded_messages.get("author", "")
                    excluded_messages["Причина"] = excluded_messages["message_id"].map(reason_map).fillna("")
                    excluded_messages["Текст"] = excluded_messages["text_clean"].fillna("").astype(str).str.slice(0, 350)
                    excluded_cols = ["Дата", "Чат", "Автор", "Причина", "Текст"]
                    excluded_select = st.dataframe(
                        excluded_messages[excluded_cols],
                        use_container_width=True,
                        height=220,
                        hide_index=True,
                        on_select="rerun",
                        selection_mode="single-row",
                        column_config={
                            "Текст": st.column_config.TextColumn(width="large"),
                        },
                    )
                    excluded_rows = get_selected_rows(excluded_select)
                    if excluded_rows:
                        restored_row = excluded_messages.iloc[excluded_rows[0]]
                        st.write(restored_row.get("text_clean", ""))
                        if st.button("Вернуть сообщение в инфоповод", key=f"restore_{event_id}_{restored_row['message_id']}"):
                            restore_message_relevance(conn, event_id, restored_row["message_id"])
                            st.success("Сообщение возвращено в инфоповод.")
                            st.cache_data.clear()
                            st.rerun()

            st.markdown("#### Массовое объединение инфоповодов")
            st.caption(
                "Можно выбрать сразу несколько инфоповодов и объединить их за одно действие. "
                "Если целевой инфоповод — текущий, выбранные темы будут присоединены к нему. "
                "Если выбран другой целевой инфоповод, текущая тема и дополнительные выбранные темы будут присоединены к нему."
            )

            def visible_event_source_ids(row_or_id) -> set[str]:
                """Return all internal source event ids hidden behind a visible/aggregated event row."""
                result: set[str] = set()
                if isinstance(row_or_id, pd.Series):
                    row = row_or_id
                else:
                    row_match = events[events["event_id"].astype(str) == str(row_or_id)]
                    row = row_match.iloc[0] if len(row_match) else pd.Series({"event_id": str(row_or_id)})

                raw = str(row.get("source_event_ids", "") or "")
                for item in raw.split("|"):
                    item = item.strip()
                    if item:
                        result.add(item)
                event_id_value = str(row.get("event_id", "") or "").strip()
                if event_id_value:
                    result.add(event_id_value)
                return result

            def make_event_label(row: pd.Series, prefix: str = "") -> str:
                title_part = str(row.get("event_title", "") or "Без названия")[:120]
                msg_count = int(row.get("message_count", 0) or 0)
                period_part = format_period(row)
                suffix = f" · {msg_count} сообщ."
                if period_part:
                    suffix += f" · {period_part}"
                return f"{prefix}{title_part}{suffix}"

            current_visible_id = str(event_id)
            current_internal_ids = visible_event_source_ids(ev)
            event_rows_for_merge = events.copy()
            event_rows_for_merge["event_id"] = event_rows_for_merge["event_id"].astype(str)
            event_rows_for_merge["search_text"] = event_rows_for_merge["event_title"].fillna("").astype(str).str.lower()

            target_search = st.text_input(
                "Поиск целевого инфоповода",
                value="",
                placeholder="Начните вводить название темы, если список большой",
                key=f"merge_target_search_{event_id}",
            )

            target_candidates = event_rows_for_merge.copy()
            if target_search.strip():
                q = target_search.strip().lower().replace("ё", "е")
                title_norm = target_candidates["search_text"].str.replace("ё", "е", regex=False)
                target_candidates = target_candidates[title_norm.str.contains(q, regex=False, na=False)]

            # Always keep the current event available as a target, even when the search filter hides it.
            current_target_row = event_rows_for_merge[event_rows_for_merge["event_id"] == current_visible_id]
            if len(current_target_row) and current_visible_id not in set(target_candidates["event_id"].astype(str)):
                target_candidates = pd.concat([current_target_row, target_candidates], ignore_index=True)

            target_candidates = target_candidates.sort_values(["message_count", "event_title"], ascending=[False, True])
            target_label_map = {}
            for _, candidate_row in target_candidates.iterrows():
                candidate_id = str(candidate_row.get("event_id", ""))
                prefix = "Текущий: " if candidate_id == current_visible_id else ""
                target_label_map[candidate_id] = make_event_label(candidate_row, prefix=prefix)

            target_ids = list(target_label_map.keys())
            if current_visible_id in target_ids:
                target_ids = [current_visible_id] + [x for x in target_ids if x != current_visible_id]

            selected_target_id = st.selectbox(
                "Целевой инфоповод",
                options=target_ids,
                format_func=lambda x: target_label_map.get(str(x), "— выберите инфоповод —"),
                key=f"merge_target_id_{event_id}",
            ) if target_ids else ""

            if not selected_target_id:
                st.info("Подходящих инфоповодов для объединения не найдено. Попробуйте изменить поиск.")

            source_search = st.text_input(
                "Поиск инфоповодов для объединения",
                value="",
                placeholder="Найдите темы, которые нужно присоединить к целевой",
                key=f"merge_sources_search_{event_id}",
            )

            target_internal_ids = visible_event_source_ids(selected_target_id) if selected_target_id else set()
            source_candidates = event_rows_for_merge.copy()
            source_candidates = source_candidates[~source_candidates["event_id"].astype(str).isin({current_visible_id, str(selected_target_id)})]

            if source_search.strip():
                q = source_search.strip().lower().replace("ё", "е")
                title_norm = source_candidates["search_text"].str.replace("ё", "е", regex=False)
                source_candidates = source_candidates[title_norm.str.contains(q, regex=False, na=False)]

            source_candidates = source_candidates.sort_values(["message_count", "event_title"], ascending=[False, True])
            source_label_map = {
                str(row.get("event_id", "")): make_event_label(row)
                for _, row in source_candidates.iterrows()
            }

            selected_extra_source_ids = st.multiselect(
                "Дополнительные инфоповоды для объединения",
                options=list(source_label_map.keys()),
                format_func=lambda x: source_label_map.get(str(x), str(x)),
                key=f"merge_extra_sources_{event_id}",
                help="Можно выбрать несколько тем. Они будут объединены за одно действие.",
            )

            if selected_target_id == current_visible_id:
                source_visible_ids = [str(x) for x in selected_extra_source_ids]
                st.caption("Режим: выбранные инфоповоды будут присоединены к текущему.")
            else:
                source_visible_ids = [current_visible_id] + [str(x) for x in selected_extra_source_ids]
                st.caption("Режим: текущий инфоповод и выбранные дополнительные темы будут присоединены к целевому.")

            internal_sources_to_merge: set[str] = set()
            for source_visible_id in source_visible_ids:
                internal_sources_to_merge.update(visible_event_source_ids(source_visible_id))
            internal_sources_to_merge = {
                source_id for source_id in internal_sources_to_merge
                if source_id and source_id not in target_internal_ids and source_id != str(selected_target_id)
            }

            if selected_target_id:
                st.caption(
                    f"Цель: {target_label_map.get(str(selected_target_id), selected_target_id)}. "
                    f"К объединению выбрано видимых тем: {len(source_visible_ids)}, внутренних событий/волн: {len(internal_sources_to_merge)}."
                )

            reason = st.text_input("Причина объединения", value="", key=f"merge_reason_{event_id}")
            merge_disabled = not bool(selected_target_id) or len(internal_sources_to_merge) == 0
            if st.button("Объединить выбранные", disabled=merge_disabled, key=f"merge_button_{event_id}"):
                target_id = str(selected_target_id)
                try:
                    merged_count = 0
                    for source_id in sorted(internal_sources_to_merge):
                        if source_id == target_id:
                            continue
                        merge_events(conn, source_event_id=source_id, target_event_id=target_id, reason=reason)
                        merged_count += 1
                    st.success(f"Инфоповоды объединены: перенесено внутренних событий/волн: {merged_count}.")
                    st.cache_data.clear()
                    st.rerun()
                except Exception as e:
                    st.error(str(e))


def show_message_search(messages: pd.DataFrame, events: pd.DataFrame, conn):
    st.subheader("Поиск по сообщениям")

    col1, col2, col3 = st.columns(3)
    q = col1.text_input("Текст")
    tag = col2.text_input("Тег")
    chat = col3.text_input("Чат")

    filtered = messages[~messages.get("message_hidden", pd.Series([False] * len(messages))).astype(bool)].copy()

    if q:
        filtered = filtered[filtered["text_clean"].fillna("").str.lower().str.contains(q.lower(), regex=False)]
    if tag and "tags" in filtered.columns:
        filtered = filtered[filtered["tags"].fillna("").str.lower().str.contains(tag.lower(), regex=False)]
    if chat and "chat_title" in filtered.columns:
        filtered = filtered[filtered["chat_title"].fillna("").str.lower().str.contains(chat.lower(), regex=False)]

    filtered = filtered.sort_values("datetime", ascending=False).head(500)
    filtered["Текст"] = filtered["text_clean"].fillna("").astype(str).str.slice(0, 350)
    filtered["Дата"] = format_date_series(filtered["datetime"])
    filtered["Чат"] = filtered.get("chat_title", "")
    filtered["Автор"] = filtered.get("author", "")
    event_title_map = events.set_index("event_id")["event_title"].to_dict() if len(events) else {}
    filtered["Инфоповод"] = filtered.get("final_event_id", "").map(event_title_map).fillna("")
    filtered["Теги"] = filtered.get("tags", "")
    filtered["Тональность"] = filtered.get("sentiment", "")
    filtered["Ссылка"] = filtered.get("message_link", "")
    cols = ["Дата", "Чат", "Автор", "Теги", "Тональность", "Инфоповод", "Текст", "Ссылка"]
    cols = [c for c in cols if c in filtered.columns]

    st.dataframe(
        filtered[cols],
        use_container_width=True,
        height=640,
        hide_index=True,
        column_config={
            "Ссылка": st.column_config.LinkColumn("Ссылка"),
            "Текст": st.column_config.TextColumn(width="large"),
        },
    )



def safe_upload_name(filename: str) -> str:
    source = Path(filename or "uploaded.csv")
    name = re.sub(r"[^0-9A-Za-zА-Яа-я_. -]+", "_", source.stem).strip("._-") or "uploaded"
    suffix = source.suffix.lower()
    if suffix not in {".csv", ".xlsx", ".xls", ".xlsm"}:
        suffix = ".csv"
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{stamp}_{name}{suffix}"


def read_manifest(data_dir: Path) -> dict:
    path = data_dir / "manifest.json"
    if not path.exists():
        return {}
    try:
        import json
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}



@st.cache_data(show_spinner=False, ttl=60)
def load_generated_tables_remote(period_ids: tuple[str, ...]):
    return load_generated_tables_from_supabase(list(period_ids))


def render_period_selector() -> list[str]:
    """Render Supabase period selector and return selected period IDs."""
    periods = list_periods()
    if periods.empty:
        st.sidebar.info("В Supabase пока нет сохраненных периодов. Загрузите файл в разделе «Загрузка файла».")
        return []

    def period_label(row) -> str:
        name = str(row.get("period_name", "") or row.get("period_id", ""))
        date_from = format_date(row.get("date_from")) if "date_from" in row else ""
        date_to = format_date(row.get("date_to")) if "date_to" in row else ""
        dates = f" · {date_from}–{date_to}" if date_from or date_to else ""
        return f"{name}{dates}"

    labels = {str(row["period_id"]): period_label(row) for _, row in periods.iterrows()}
    options = list(labels.keys())
    default = options[:1]
    selected = st.sidebar.multiselect(
        "Периоды",
        options=options,
        default=default,
        format_func=lambda x: labels.get(x, x),
        help="Можно выбрать один период или несколько. При выборе нескольких периодов дашборд объединит данные.",
    )
    if not selected:
        st.sidebar.warning("Выберите хотя бы один период.")
    return selected
def render_period_history_manager(persistent_enabled: bool, upload_dir: Path):
    """Show and edit uploaded periods/files history."""
    st.markdown("### История загруженных файлов")

    if not persistent_enabled:
        st.info("Supabase не настроен. В локальном режиме можно только посмотреть файлы в папке загрузок.")
        upload_dir.mkdir(parents=True, exist_ok=True)
        files = sorted([p for p in upload_dir.iterdir() if p.suffix.lower() in {".csv", ".xlsx", ".xls", ".xlsm"}], reverse=True)
        if not files:
            st.write("Пока нет загруженных файлов.")
            return
        history = pd.DataFrame({
            "Файл": [f.name for f in files],
            "Размер, КБ": [round(f.stat().st_size / 1024, 1) for f in files],
            "Дата загрузки": [datetime.fromtimestamp(f.stat().st_mtime).strftime("%d.%m.%Y") for f in files],
        })
        st.dataframe(history, use_container_width=True, hide_index=True)
        if st.button("Очистить локальную историю загруженных файлов"):
            for f in files:
                f.unlink(missing_ok=True)
            st.success("Локальная история загрузок очищена.")
            st.rerun()
        return

    try:
        periods = list_periods(include_inactive=True)
    except Exception as e:
        st.warning(f"Не удалось получить историю периодов из Supabase: {e}")
        return

    if periods.empty:
        st.write("Пока нет сохраненных периодов в Supabase.")
        return

    view = periods.copy().reset_index(drop=True)
    for col in ["date_from", "date_to", "uploaded_at"]:
        if col in view.columns:
            view[col] = pd.to_datetime(view[col], errors="coerce").dt.strftime("%d.%m.%Y").fillna("")
    if "manifest" in view.columns:
        view["period_note"] = view["manifest"].apply(period_note_from_manifest)
    else:
        view["period_note"] = ""

    display_cols = [c for c in ["period_name", "date_from", "date_to", "source_filename", "uploaded_at", "status", "period_note"] if c in view.columns]
    display = view[display_cols].rename(columns={
        "period_name": "Название периода",
        "date_from": "Начало",
        "date_to": "Конец",
        "source_filename": "Файл",
        "uploaded_at": "Загружено",
        "status": "Статус",
        "period_note": "Комментарий",
    })

    event = st.dataframe(
        display,
        use_container_width=True,
        hide_index=True,
        height=320,
        selection_mode="single-row",
        on_select="rerun",
    )
    rows = get_selected_rows(event)
    if not rows:
        st.caption("Выберите строку в таблице, чтобы изменить название периода, даты, статус или комментарий.")
        return

    selected = periods.reset_index(drop=True).iloc[rows[0]].to_dict()
    period_id = str(selected.get("period_id", ""))
    st.markdown("#### Редактирование выбранного периода")
    st.caption(f"Внутренний ID: `{period_id}`")

    current_manifest = selected.get("manifest") or {}
    if not isinstance(current_manifest, dict):
        current_manifest = {}

    date_from_value = parse_date_value(selected.get("date_from")) or datetime.now().date()
    date_to_value = parse_date_value(selected.get("date_to")) or date_from_value
    status_options = ["active", "hidden", "archived"]
    current_status = str(selected.get("status") or "active")
    if current_status not in status_options:
        status_options.append(current_status)

    with st.form(f"period_edit_form_{period_id}"):
        c1, c2 = st.columns(2)
        new_name = c1.text_input("Название периода", value=str(selected.get("period_name") or ""))
        new_source = c2.text_input("Название исходного файла", value=str(selected.get("source_filename") or ""))
        d1, d2 = st.columns(2)
        new_date_from = d1.date_input(
            "Дата начала",
            value=date_from_value,
            format="DD.MM.YYYY",
            help="Дата хранится в Supabase как YYYY-MM-DD, а в интерфейсе отображается как ДД.ММ.ГГГГ.",
        )
        new_date_to = d2.date_input(
            "Дата окончания",
            value=date_to_value,
            format="DD.MM.YYYY",
            help="Дата хранится в Supabase как YYYY-MM-DD, а в интерфейсе отображается как ДД.ММ.ГГГГ.",
        )
        new_status = st.selectbox(
            "Статус периода",
            status_options,
            index=status_options.index(current_status),
            help="active — показывать в фильтре периодов; hidden/archived — скрыть из основного фильтра, но оставить в истории.",
        )
        new_note = st.text_area("Комментарий к периоду", value=str(current_manifest.get("period_note", "") or ""), height=90)
        submitted = st.form_submit_button("Сохранить изменения", type="primary")

    if submitted:
        parsed_from = date_to_iso(new_date_from)
        parsed_to = date_to_iso(new_date_to)
        if not parsed_from:
            st.error("Не удалось распознать дату начала.")
            return
        if not parsed_to:
            st.error("Не удалось распознать дату окончания.")
            return
        if parsed_from > parsed_to:
            st.error("Дата начала не может быть позже даты окончания.")
            return
        try:
            update_period_metadata(
                period_id,
                period_name=new_name,
                date_from=parsed_from,
                date_to=parsed_to,
                source_filename=new_source,
                status=new_status,
                manifest_updates={"period_note": new_note},
            )
            st.cache_data.clear()
            st.success("Данные периода обновлены.")
            st.rerun()
        except Exception as e:
            st.error("Не удалось сохранить изменения периода.")
            st.exception(e)

    c1, c2, c3 = st.columns(3)
    if c1.button("Скрыть из фильтра", key=f"hide_period_{period_id}"):
        set_period_status(period_id, "hidden")
        st.cache_data.clear()
        st.success("Период скрыт из основного фильтра.")
        st.rerun()
    if c2.button("Вернуть в фильтр", key=f"activate_period_{period_id}"):
        set_period_status(period_id, "active")
        st.cache_data.clear()
        st.success("Период снова активен.")
        st.rerun()

    with st.expander("Удаление периода", expanded=False):
        st.warning("Полное удаление уберет период и его обработанные строки из Supabase. Это действие нельзя отменить из дашборда.")
        confirm_delete = st.checkbox("Я понимаю, что период будет удален полностью", key=f"confirm_delete_period_{period_id}")
        if st.button("Удалить период полностью", disabled=not confirm_delete, key=f"hard_delete_period_{period_id}"):
            try:
                delete_period(period_id, hard=True)
                st.cache_data.clear()
                st.success("Период удален.")
                st.rerun()
            except Exception as e:
                st.error("Не удалось удалить период.")
                st.exception(e)


def show_upload_page(data_dir: Path, upload_dir: Path):
    st.subheader("Загрузка файла нового периода")
    st.write(
        "Загрузите новый CSV или Excel-файл из Brand Analytics, Медиалогии или другой системы. "
        "Дашборд приведет данные к единому формату и пересоберет сообщения, обсуждения и инфоповоды."
    )

    persistent_enabled = supabase_configured()
    selected_period_ids: list[str] = []
    if persistent_enabled:
        st.success("Постоянное хранение включено: исходные файлы и обработанные периоды будут сохраняться в Supabase.")
    else:
        st.warning(
            "Supabase не настроен. Загруженные файлы и пересчитанные таблицы сохранятся только в текущем runtime Streamlit Cloud "
            "и после redeploy могут сброситься к версии из GitHub."
        )

    manifest = read_manifest(data_dir)
    if manifest:
        c1, c2, c3 = st.columns(3)
        c1.metric("Текущих сообщений", f"{int(manifest.get('rows_messages', 0)):,}".replace(",", " "))
        c2.metric("Текущих обсуждений", f"{int(manifest.get('rows_discussions', 0)):,}".replace(",", " "))
        c3.metric("Текущих инфоповодов", f"{int(manifest.get('rows_events', 0)):,}".replace(",", " "))

    default_period_name = datetime.now().strftime("%d.%m.%Y")
    period_name = st.text_input(
        "Название периода",
        value=default_period_name,
        help="Например: 24.04.2026–30.04.2026. Так период будет отображаться в фильтре.",
    )

    mode = st.radio(
        "Режим загрузки",
        [
            "Заменить текущую выборку новым файлом",
            "Добавить файл в историю загрузок и пересобрать все загруженные периоды",
        ],
        help=(
            "В режиме добавления дашборд объединит все CSV/Excel-файлы из папки data/uploads. "
            "Если исходный файл старого периода не был загружен в историю, он не попадет в пересборку."
        ),
    )

    uploaded = st.file_uploader("Файл нового периода", type=["csv", "xlsx", "xls", "xlsm"])

    with st.expander("Параметры алгоритма", expanded=False):
        col1, col2 = st.columns(2)
        window_minutes = col1.number_input("Окно обсуждения, минут", min_value=10, max_value=240, value=60, step=5)
        similarity_threshold = col2.slider("Порог похожести", min_value=0.10, max_value=0.60, value=0.28, step=0.01)
        event_gap_hours = col1.number_input("Разрыв между волнами, часов", min_value=0.5, max_value=24.0, value=3.0, step=0.5)
        event_window_hours = col2.number_input("Максимальная длина волны, часов", min_value=2.0, max_value=72.0, value=16.0, step=2.0)
        cluster_method = st.selectbox("Метод кластеризации", ["tfidf", "none"], index=0)

    col_a, col_b = st.columns([1, 2])
    run = col_a.button("Загрузить и пересобрать", type="primary", disabled=uploaded is None)

    if run and uploaded is not None:
        upload_dir.mkdir(parents=True, exist_ok=True)
        data_dir.mkdir(parents=True, exist_ok=True)

        saved_path = upload_dir / safe_upload_name(uploaded.name)
        saved_path.write_bytes(uploaded.getbuffer())

        try:
            with st.spinner("Читаю файл и пересобираю инфоповоды…"):
                if persistent_enabled or mode.startswith("Заменить"):
                    manifest = run_preprocess(
                        input_path=saved_path,
                        output=data_dir,
                        window_minutes=int(window_minutes),
                        cluster_method=cluster_method,
                        similarity_threshold=float(similarity_threshold),
                        event_gap_hours=float(event_gap_hours),
                        event_window_hours=float(event_window_hours),
                    )
                else:
                    csv_paths = sorted([p for p in upload_dir.iterdir() if p.suffix.lower() in {".csv", ".xlsx", ".xls", ".xlsm"}])
                    raw_parts = []
                    used_files = []
                    for path in csv_paths:
                        try:
                            raw_parts.append(read_source_csv(path))
                            used_files.append(path.name)
                        except Exception as e:
                            st.warning(f"Файл {path.name} пропущен: {e}")
                    if not raw_parts:
                        st.error("Не удалось прочитать ни один файл.")
                        return
                    combined = pd.concat(raw_parts, ignore_index=True)
                    manifest = run_preprocess_from_dataframe(
                        raw=combined,
                        output=data_dir,
                        source_file="; ".join(used_files),
                        window_minutes=int(window_minutes),
                        cluster_method=cluster_method,
                        similarity_threshold=float(similarity_threshold),
                        event_gap_hours=float(event_gap_hours),
                        event_window_hours=float(event_window_hours),
                    )


            if persistent_enabled:
                period_title = str(period_name or uploaded.name).strip() or uploaded.name
                period_id = make_period_id(period_title, uploaded.name)
                with st.spinner("Сохраняю период в Supabase…"):
                    save_processed_tables_from_dir(
                        data_dir,
                        period_id=period_id,
                        period_name=period_title,
                        source_filename=uploaded.name,
                        manifest=manifest,
                        replace=True,
                    )
                    try:
                        save_uploaded_csv_to_storage(period_id, uploaded.name, bytes(uploaded.getbuffer()))
                    except Exception as storage_error:
                        st.warning(f"Обработанные данные сохранены в Supabase, но исходный файл не удалось сохранить в Storage: {storage_error}")
                st.success(f"Период сохранен в Supabase: {period_title}")

            st.cache_data.clear()
            st.success(
                "Данные пересобраны: "
                f"{manifest.get('rows_messages', 0)} сообщений, "
                f"{manifest.get('rows_discussions', 0)} обсуждений, "
                f"{manifest.get('rows_events', 0)} инфоповодов."
            )
            st.info("Перейдите в раздел «Инфоповоды» или обновите страницу, чтобы увидеть новую выборку.")
            if st.button("Открыть обновленные инфоповоды"):
                st.rerun()
        except Exception as e:
            st.error("Не удалось обработать файл.")
            st.exception(e)

    with st.expander("История и редактирование загруженных файлов", expanded=False):
        render_period_history_manager(persistent_enabled, upload_dir)


def main():
    args = parse_args()

    st.set_page_config(
        page_title="Дайджест водительских чатов",
        layout="wide",
        initial_sidebar_state="expanded",
    )

    st.title("Дайджест водительских чатов")

    data_dir = Path(args.data_dir)
    db_path = Path(args.db_path)
    upload_dir = Path(args.upload_dir)
    persistent_enabled = supabase_configured()
    selected_period_ids: list[str] = []

    if persistent_enabled:
        st.sidebar.success("Хранилище: Supabase")
    else:
        st.sidebar.info("Хранилище: локальные файлы")

    can_edit = render_admin_mode()
    if can_edit:
        st.caption("Версия 3.6: добавлено массовое объединение нескольких инфоповодов")
    pages = ["Инфоповоды", "Поиск сообщений"] + (["Загрузка файла"] if can_edit else [])
    page = st.sidebar.radio("Раздел", pages, label_visibility="collapsed")

    if page == "Загрузка файла":
        show_upload_page(data_dir, upload_dir)
        return

    conn = connect(db_path)

    if persistent_enabled:
        try:
            selected_period_ids = render_period_selector()
            if not selected_period_ids:
                st.info("Пока нет выбранных периодов. Откройте «Загрузка файла» и сохраните первый период в Supabase.")
                st.stop()
            events_raw, discussions, messages, discussion_messages, event_discussions = load_generated_tables_remote(tuple(selected_period_ids))
        except Exception as e:
            st.error("Не удалось загрузить данные из Supabase.")
            st.exception(e)
            st.stop()
    else:
        if not data_dir.exists():
            st.error(f"Папка с обработанными данными не найдена: {data_dir}")
            st.info("Откройте раздел «Загрузка файла» и загрузите исходный файл для первой сборки дашборда.")
            st.stop()
        events_raw, discussions, messages, discussion_messages, event_discussions = load_generated_tables(str(data_dir))

    overrides = get_event_overrides(conn)
    merges = get_event_merges(conn)
    msg_overrides = get_message_overrides(conn)
    msg_exclusions = get_message_exclusions(conn)
    manual_events = get_manual_events(conn)
    msg_topic_overrides = get_message_topic_overrides(conn)

    events, enriched_messages = apply_manual_edits(
        events_raw,
        messages,
        discussion_messages,
        event_discussions,
        overrides,
        merges,
        msg_overrides,
        msg_exclusions,
        manual_events,
        msg_topic_overrides,
    )

    if page == "Поиск сообщений":
        show_message_search(enriched_messages, events, conn)
        return

    if can_edit:
        show_manual_event_creator(conn)
    else:
        st.sidebar.caption("Загрузка файлов и ручная модерация доступны после входа администратора.")

    if persistent_enabled and len(selected_period_ids) >= 2:
        render_period_comparison(enriched_messages, selected_period_ids)

    render_dashboard_summary(events, enriched_messages, selected_period_ids, conn, can_edit)
    render_digest_export(events, enriched_messages, selected_period_ids, conn)

    filtered_events, word_query, word_matches = apply_filters(events, enriched_messages)
    filtered_messages = messages_for_events(enriched_messages, filtered_events)
    show_kpis(filtered_events, filtered_messages)
    if word_query:
        word_message_results_table(word_matches, events, word_query)
    selected_event_id = event_table(filtered_events)
    if selected_event_id:
        show_event_card(selected_event_id, events, enriched_messages, conn, can_edit=can_edit)


if __name__ == "__main__":
    main()
